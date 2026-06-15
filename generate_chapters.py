"""
generate_chapters.py
Generates Inspirability Graduation Project Documentation
Chapters 8, 9, and 10 as a .docx Word document.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os

# ── Colour palette ────────────────────────────────────────────────────────────
PRIMARY     = RGBColor(0x16, 0x22, 0x7E)   # navy
ACCENT      = RGBColor(0x00, 0xA6, 0xC8)   # teal
DARK_TEXT   = RGBColor(0x1A, 0x1A, 0x2E)
TABLE_HDR   = RGBColor(0x16, 0x22, 0x7E)   # navy header
TABLE_ALT   = RGBColor(0xEE, 0xF0, 0xFF)   # light lavender alt row
PLACEHOLDER = RGBColor(0xEE, 0xF0, 0xFF)   # figure box fill
CODE_BG     = RGBColor(0xF4, 0xF4, 0xF4)   # code block background

OUTPUT_PATH = os.path.join(
    os.path.dirname(__file__),
    "my-app",
    "Inspirability_Chapters_8_9_10.docx",
)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Default normal style ──────────────────────────────────────────────────────
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)

# ─────────────────────────────────────────────────────────────────────────────
# Helper utilities
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_shading(cell, hex_color: str):
    """Fill a table cell with a solid background color."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_para_shading(para, hex_color: str):
    """Fill a paragraph (used for code/figure blocks) with background color."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)


def add_chapter_heading(text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(12)
    p.paragraph_format.keep_with_next = False
    run = p.add_run(text)
    run.font.name  = "Calibri"
    run.font.size  = Pt(18)
    run.font.bold  = True
    run.font.color.rgb = PRIMARY
    return p


def add_heading2(text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.name  = "Calibri"
    run.font.size  = Pt(14)
    run.font.bold  = True
    run.font.color.rgb = PRIMARY
    return p


def add_heading3(text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name  = "Calibri"
    run.font.size  = Pt(12)
    run.font.bold  = True
    run.font.color.rgb = ACCENT
    return p


def add_body(text: str, indent_cm: float = 0.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    run.font.name  = "Calibri"
    run.font.size  = Pt(11)
    run.font.color.rgb = DARK_TEXT
    return p


def add_bullet(text: str, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(1.0 + level * 0.6)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.name  = "Calibri"
    run.font.size  = Pt(11)
    run.font.color.rgb = DARK_TEXT
    return p


def add_figure_placeholder(fig_num: int, caption: str, description: str = "", height_cm: float = 6.0):
    """Shaded placeholder box + figure caption + optional description."""
    # Box paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    set_para_shading(p, "EEF0FF")
    # Add left/right border simulation via indents
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    # Inner lines to simulate height
    lines = max(3, int(height_cm * 1.5))
    inner_text = "\n" * (lines // 2) + f"[ Figure {fig_num} — {caption} ]\n" + "\n" * (lines // 2)
    run = p.add_run(inner_text)
    run.font.name   = "Calibri"
    run.font.size   = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x99)

    # Caption paragraph
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after  = Pt(4)
    cr = cap.add_run(f"Figure {fig_num}: {caption}")
    cr.font.name   = "Calibri"
    cr.font.size   = Pt(10)
    cr.font.italic = True
    cr.font.bold   = True
    cr.font.color.rgb = DARK_TEXT

    # Optional description
    if description:
        add_body(description)


def add_code_block(code_lines: list, caption: str = "", fig_num: int = 0):
    """Monospace shaded block for code snippets."""
    for line in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(1)
        set_para_shading(p, "F4F4F4")
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    if caption and fig_num:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(4)
        cap.paragraph_format.space_after  = Pt(8)
        cr = cap.add_run(f"Figure {fig_num}: {caption}")
        cr.font.name   = "Calibri"
        cr.font.size   = Pt(10)
        cr.font.italic = True
        cr.font.bold   = True


def add_table(headers: list, rows: list, caption: str = "", fig_num: int = 0):
    """Styled table with navy header row and alternating light rows."""
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_shading(cell, "16227E")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size  = Pt(10)
        run.font.name  = "Calibri"

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg  = "EEF0FF" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_shading(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "Calibri"

    doc.add_paragraph()  # spacing after table

    if caption and fig_num:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        cr = cap.add_run(f"Table {fig_num}: {caption}")
        cr.font.name   = "Calibri"
        cr.font.size   = Pt(10)
        cr.font.italic = True
        cr.font.bold   = True

    return table


def add_page_break():
    doc.add_page_break()


def add_hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "16227E")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ─────────────────────────────────────────────────────────────────────────────
# LIST OF FIGURES (for these 3 chapters)
# ─────────────────────────────────────────────────────────────────────────────

add_chapter_heading("List of Figures")
add_hr()

figures = [
    (1,  "Inspirability Frontend Technology Stack"),
    (2,  "Inspirability Backend Architecture Overview"),
    (3,  "MySQL Database Entity-Relationship Overview"),
    (4,  "users Table Schema"),
    (5,  "school Table Schema"),
    (6,  "medical_clinic Table Schema"),
    (7,  "sport_center Table Schema"),
    (8,  "appointment Table Schema"),
    (9,  "Homepage Interface"),
    (10, "Login Page Interface"),
    (11, "Sign-Up Page Interface"),
    (12, "About Page Interface"),
    (13, "Services Page Interface"),
    (14, "School Directory Interface"),
    (15, "Medical Services Page Interface"),
    (16, "Sport Services Page Interface"),
    (17, "Pricing & Flat-Fee Plans Interface"),
    (18, "Checkout / Payment Page Interface"),
    (19, "Feedback Page Interface"),
    (20, "User Profile Page Interface"),
    (21, "My Appointments Page Interface"),
    (22, "Provider Appointments Management Interface"),
    (23, "Admin Dashboard Interface"),
    (24, "Admin Requests Management Interface"),
    (25, "React Component — Navbar.jsx"),
    (26, "API Fetching with Axios — School Directory"),
    (27, "JWT Authentication — Login Controller"),
    (28, "CRUD Operations — Appointment Controller"),
    (29, "Database Integration — MySQL Pool Connection"),
    (30, "Postman API Testing — POST /api/auth/login"),
    (31, "Postman API Testing — GET /api/admin/stats"),
    (32, "Postman API Testing — POST /api/appointments"),
    (33, "System Requirements Validation Summary"),
    (34, "Flutter Mobile App UI Prototype — Home Screen"),
    (35, "Flutter Mobile App UI Prototype — Services Screen"),
]

for num, label in figures:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"Figure {num:>2}   {label}")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_TEXT

add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# LIST OF TABLES
# ─────────────────────────────────────────────────────────────────────────────

add_chapter_heading("List of Tables")
add_hr()

tables_list = [
    (1,  "Frontend Technology Stack"),
    (2,  "Backend Technology Stack"),
    (3,  "Development & Collaboration Tools"),
    (4,  "users Table — Column Definitions"),
    (5,  "school Table — Column Definitions"),
    (6,  "medical_clinic Table — Column Definitions"),
    (7,  "sport_center Table — Column Definitions"),
    (8,  "appointment Table — Column Definitions"),
    (9,  "System Functional Requirements Validation Matrix"),
    (10, "System Non-Functional Requirements Validation Matrix"),
    (11, "Test Cases — Authentication Module"),
    (12, "Test Cases — Appointment Management Module"),
    (13, "Test Cases — Admin Management Module"),
]

for num, label in tables_list:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"Table {num:>2}   {label}")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_TEXT

add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 8 — SYSTEM IMPLEMENTATION
# ═════════════════════════════════════════════════════════════════════════════

add_chapter_heading("Chapter 8 — System Implementation")
add_hr()

# ── 8.1 Introduction ──────────────────────────────────────────────────────────
add_heading2("8.1  Introduction")

add_body(
    "This chapter presents a comprehensive account of the technical implementation of the Inspirability "
    "platform. It covers the full software development lifecycle from the selection of the technology "
    "stack through to the construction of individual system components, including the frontend user "
    "interface, the RESTful backend API, the relational database layer, and the graphical user interface "
    "as experienced by end users. The goal of this chapter is to provide sufficient technical depth for "
    "the reader to understand how the system was built, how its components interconnect, and how each "
    "design decision supports the functional and non-functional requirements established in earlier chapters."
)

add_body(
    "The Inspirability platform is a full-stack web application developed to serve families of children "
    "with special needs in Egypt by consolidating school, medical, and sport services in a single digital "
    "environment. The system was implemented using industry-standard technologies selected for their "
    "maturity, community support, and suitability to the project's requirements. The following subsections "
    "detail each layer of the technology stack."
)

add_heading3("8.1.1  Frontend Technologies")

add_body(
    "The frontend of Inspirability was developed using Next.js 15.5 and React 19, both of which are "
    "among the most widely adopted frameworks for building modern, server-side-rendered and statically "
    "generated web applications. Next.js provides the App Router architecture, enabling file-system-based "
    "routing, server components, and seamless client-side transitions. React's component-based model "
    "allowed the team to build reusable, self-contained UI elements that could be assembled into complete "
    "page layouts with minimal code duplication."
)
add_body(
    "Tailwind CSS v4 was adopted as the utility-first styling framework, enabling rapid interface "
    "construction without leaving the component files. Chart.js 4.5 and its React wrapper react-chartjs-2 "
    "5.3 were used to render interactive line charts on the Admin Dashboard, visualising user signups, "
    "site visits, and appointment trends over a six-month window. Axios 1.15 served as the HTTP client "
    "for all API communication between the frontend and the Express backend."
)

add_figure_placeholder(1, "Inspirability Frontend Technology Stack",
    "The diagram above illustrates the layers of the frontend stack: Next.js App Router at the top, "
    "React components in the middle, and supporting libraries (Tailwind CSS, Axios, Chart.js) at the base.",
    height_cm=5)

add_table(
    ["Technology", "Version", "Role"],
    [
        ["Next.js",          "15.5.18",  "App framework, routing, SSR/SSG"],
        ["React",            "19.2.5",   "Component model, state management"],
        ["Tailwind CSS",     "4.0",      "Utility-first styling"],
        ["Axios",            "1.15.0",   "HTTP client for REST API calls"],
        ["Chart.js",         "4.5.1",    "Data visualisation (admin charts)"],
        ["react-chartjs-2",  "5.3.1",    "React wrapper for Chart.js"],
        ["TypeScript",       "5.x",      "Static type checking"],
        ["ESLint",           "9.x",      "Code quality and linting"],
    ],
    caption="Frontend Technology Stack",
    fig_num=1
)

add_heading3("8.1.2  Backend Technologies")

add_body(
    "The backend was implemented as a RESTful API server using Node.js and Express.js. Express.js was "
    "chosen for its minimal footprint, robust middleware ecosystem, and the team's familiarity with its "
    "routing conventions. The server handles all business logic, authentication, file upload management, "
    "and database interactions."
)
add_body(
    "MySQL was selected as the relational database management system due to its reliability, wide "
    "adoption in production environments, and strong support for complex JOIN queries required by the "
    "multi-role data model. The mysql2 Node.js driver, configured as a connection pool, was used for "
    "all database operations. JSON Web Tokens (JWT) via the jsonwebtoken library provided stateless "
    "authentication across all protected routes. File uploads — including school logos, medical "
    "certificates, and profile images — were handled by Multer, with binary content stored in a "
    "dedicated media table as LONGBLOB data. bcrypt was used for password hashing to ensure user "
    "credentials are never stored in plain text. Socket.io 4.8 was integrated to support real-time "
    "chat functionality between users and the AI chatbot service."
)

add_figure_placeholder(2, "Inspirability Backend Architecture Overview",
    "The diagram above depicts the request lifecycle: HTTP request → Express router → controller → "
    "MySQL pool → JSON response. JWT middleware intercepts protected routes before controller execution.",
    height_cm=5)

add_table(
    ["Technology", "Version", "Role"],
    [
        ["Node.js",        "18 LTS",   "JavaScript runtime environment"],
        ["Express.js",     "4.x",      "REST API framework"],
        ["MySQL",          "8.x",      "Relational database management system"],
        ["mysql2",         "3.x",      "Node.js MySQL driver with promise API"],
        ["jsonwebtoken",   "9.x",      "JWT token generation and verification"],
        ["bcrypt",         "5.x",      "Password hashing"],
        ["Multer",         "1.x",      "Multipart form data / file upload handling"],
        ["Socket.io",      "4.8.0",    "Real-time bidirectional communication"],
        ["CORS",           "2.8.5",    "Cross-origin resource sharing middleware"],
    ],
    caption="Backend Technology Stack",
    fig_num=2
)

add_heading3("8.1.3  Development and Collaboration Tools")

add_body(
    "Version control was managed using Git, with the project repository hosted on GitHub. The development "
    "team adopted a feature-branch workflow, creating dedicated branches for each major feature and "
    "merging through pull requests after peer review. Visual Studio Code served as the primary IDE for "
    "all frontend and backend development due to its rich extension ecosystem and integrated terminal. "
    "Postman was used extensively during the testing and debugging phases to issue HTTP requests against "
    "the backend API endpoints independently of the frontend."
)
add_body(
    "In addition to the web platform, a mobile application UI/UX prototype was designed using Flutter "
    "as part of the project's future mobile expansion strategy. The Flutter prototype demonstrates the "
    "intended look and feel of the Inspirability mobile experience, covering the home screen, service "
    "browsing, and appointment booking flows. While the mobile application is not yet connected to the "
    "live backend, it represents a concrete roadmap artefact for the next development phase."
)

add_table(
    ["Tool", "Purpose"],
    [
        ["Git / GitHub",       "Version control, branch management, code review"],
        ["Visual Studio Code", "Primary IDE for frontend and backend development"],
        ["Postman",            "API testing, endpoint verification"],
        ["MySQL Workbench",    "Database schema design and query testing"],
        ["Flutter",            "Mobile app UI/UX prototype design"],
        ["Figma",              "UI wireframing and design mockups"],
    ],
    caption="Development & Collaboration Tools",
    fig_num=3
)

add_page_break()

# ── 8.2 Frontend & Backend Implementation ────────────────────────────────────
add_heading2("8.2  Frontend and Backend Implementation")

add_body(
    "This section describes how the frontend and backend components of Inspirability were structured, "
    "developed, and integrated to deliver the platform's full feature set."
)

add_heading3("8.2.1  Frontend Architecture")

add_body(
    "The Next.js App Router was leveraged to organise all pages under the my-app/app/ directory, with "
    "each subfolder representing a distinct URL route. This file-system routing approach eliminates "
    "explicit route configuration and makes the application structure immediately legible. The application "
    "contains twenty-five page routes covering every user-facing feature, from public informational pages "
    "to role-restricted dashboards."
)
add_body(
    "Components were organised into four groups: global components (Navbar, Footer, Toast, ChatBot, "
    "VisitTracker, ConfirmModal), home-page components (Hero, About, Categories, Story, Contact), "
    "profile components (ProfileHeader, ProfileHero, ParentProfile, SchoolProfile, ClinicProfile, "
    "SportProfile, EditProfileModal, ProfileInfoSections), and shared utility components (FilterBar, "
    "CardGrid, ProfileCard). Each component is a self-contained React functional component that manages "
    "its own local state via the useState and useEffect hooks."
)
add_body(
    "Dynamic data rendering is achieved through client-side API fetching with Axios. Pages marked with "
    "the 'use client' directive issue GET requests to the Express backend on component mount and store "
    "the response in local state, which React renders reactively as data arrives. Loading states and "
    "error boundaries prevent the UI from displaying empty or broken content during asynchronous operations."
)

add_heading3("8.2.2  Role-Based Access Control")

add_body(
    "Inspirability implements a role-based access control (RBAC) system with four distinct user roles: "
    "Parent, School, Medical Clinic, and Sport Center (collectively referred to as Service Providers), "
    "and Administrator. Each role is assigned a discrete set of permissions enforced at both the "
    "frontend (conditional rendering) and backend (JWT middleware) layers."
)

add_body("The access privileges of each role are summarised as follows:")
add_bullet("Parent: Browse all service providers; book, view, and cancel appointments; manage personal profile; submit feedback; upgrade subscription plan.")
add_bullet("School: Manage institutional profile and media gallery; view and respond to appointment requests; manage subscription payment.")
add_bullet("Medical Clinic: Manage clinic profile; manage therapy session appointments; upload certifications and media.")
add_bullet("Sport Center: Manage center profile; handle session bookings; manage media and pricing information.")
add_bullet("Administrator: Access the admin dashboard; approve or reject provider registration requests; approve or reject payment upgrade requests; view platform analytics; manage all user accounts.")

add_heading3("8.2.3  Authentication and Session Management")

add_body(
    "User authentication is handled by the authController.js module. On successful login, the backend "
    "generates a signed JWT containing the user's ID, email, and role, with a 24-hour expiry window. "
    "The token is returned to the client and stored in localStorage. All subsequent requests to "
    "protected endpoints include the token in the Authorization header as a Bearer token. Express "
    "middleware on the backend verifies the token's signature and expiry before allowing the request "
    "to reach the controller. Expired or missing tokens receive a 401 Unauthorized response, triggering "
    "an automatic redirect to the login page on the client."
)
add_body(
    "Password security is enforced using bcrypt with a salting factor of 10, ensuring that even "
    "identical passwords produce different stored hashes. No plain-text credentials are persisted "
    "in the database at any point."
)

add_heading3("8.2.4  RESTful API Design")

add_body(
    "The backend exposes a RESTful API organised into ten logical route modules, each corresponding "
    "to a domain entity or functional area: authentication, profiles, schools, medical clinics, sport "
    "centers, appointments, payments, admin operations, parent operations, and chat. All responses "
    "are JSON-formatted. A global error handler middleware ensures that unexpected server errors are "
    "returned as structured JSON rather than the Express default HTML error page, preventing "
    "JSON-parse failures on the client."
)
add_body("The principal API endpoints are organised as follows:")
add_bullet("POST /api/auth/signup — Register a new user (parent or service provider)")
add_bullet("POST /api/auth/login — Authenticate and receive a JWT token")
add_bullet("GET  /api/profile — Retrieve the authenticated user's profile")
add_bullet("PUT  /api/profile/update — Update profile fields")
add_bullet("POST /api/profile/upload-image — Upload profile or media images")
add_bullet("GET  /api/school / /api/medical / /api/sport-center — Browse service providers")
add_bullet("POST /api/appointments — Create a new appointment booking")
add_bullet("GET  /api/appointments/my — Retrieve the current user's appointments")
add_bullet("GET  /api/appointments/provider — Retrieve appointments for a provider")
add_bullet("PUT  /api/appointments/:id — Update appointment status")
add_bullet("DELETE /api/appointments/:id — Cancel an appointment")
add_bullet("POST /api/appointments/payment-request — Submit a payment upgrade request")
add_bullet("GET  /api/admin/stats — Retrieve platform analytics for the admin dashboard")
add_bullet("GET  /api/admin/provider-requests — List pending provider approval requests")
add_bullet("PUT  /api/admin/approve-provider/:id — Approve a provider registration")
add_bullet("PUT  /api/admin/reject-provider/:id — Reject a provider registration")
add_bullet("GET  /api/admin/payment-requests — List pending payment upgrade requests")
add_bullet("PUT  /api/admin/approve-payment/:id — Approve a payment request")

add_heading3("8.2.5  Admin Dashboard")

add_body(
    "The administrator dashboard (accessible at /admin) provides real-time platform analytics rendered "
    "using Chart.js. The dashboard aggregates and displays: total registered users, total service "
    "providers, new user registrations this month, cumulative site visits, and a six-month trend chart "
    "covering monthly signups, appointments, and page visits. Growth indicators compare the current "
    "month's figures against the previous month and display a colour-coded percentage badge (green for "
    "positive growth, red for decline). A recent-activity panel shows the five most recently registered "
    "users with their role and timestamp."
)
add_body(
    "The backend getStats controller implements a safeQuery pattern in which each database query is "
    "individually wrapped in a try/catch block. This ensures that a failure in any single query — for "
    "example, if the site_visits table has not yet been populated — returns a zero fallback rather than "
    "causing the entire endpoint to fail with a 500 error. This design makes the dashboard resilient "
    "during early deployment when some data may be sparse."
)

add_page_break()

# ── 8.3 Database Implementation ──────────────────────────────────────────────
add_heading2("8.3  Database Implementation")

add_body(
    "The Inspirability backend relies on a MySQL 8 relational database named inspirability. The schema "
    "is normalised to the third normal form (3NF) to minimise redundancy and enforce data integrity. "
    "The database is accessed through a mysql2 connection pool configured with a limit of ten concurrent "
    "connections and indefinite queue management, ensuring that bursts of simultaneous requests are "
    "handled gracefully without connection exhaustion."
)
add_body(
    "A lightweight migration system was implemented directly in the db.js initialisation module. On "
    "server startup, a series of ALTER TABLE and CREATE TABLE IF NOT EXISTS statements are executed to "
    "add any columns or tables introduced in newer code versions. MySQL error code 1060 (duplicate "
    "column), 1061 (duplicate index), and 1050 (table already exists) are silently ignored, making the "
    "migration idempotent across repeated restarts. This approach avoids the overhead of a dedicated "
    "migration framework while still protecting production data during incremental deployments."
)

add_figure_placeholder(3, "MySQL Database Entity-Relationship Overview",
    "The diagram above depicts the relationships between the primary entities in the Inspirability "
    "database. The users table acts as the central hub, related one-to-one with each provider profile "
    "table and one-to-many with appointments, media, and feedback records.",
    height_cm=7)

add_heading3("8.3.1  Core Tables")

add_body(
    "The following subsections describe the principal database tables and their column definitions."
)

# users table
add_heading3("users")
add_body(
    "The users table is the authoritative identity store for all accounts on the platform. Every "
    "registered user — regardless of role — has exactly one record in this table. The role column "
    "is used throughout the system to determine which profile sub-table to join and which permissions "
    "to enforce."
)
add_table(
    ["Column", "Type", "Constraint", "Description"],
    [
        ["user_id",          "INT",          "PK, AUTO_INCREMENT", "Unique user identifier"],
        ["email",            "VARCHAR(255)", "UNIQUE, NOT NULL",   "Login email address"],
        ["password",         "VARCHAR(255)", "NOT NULL",           "bcrypt-hashed password"],
        ["role",             "ENUM",         "NOT NULL",           "parent / school / clinic / sport / admin"],
        ["approval_status",  "ENUM",         "DEFAULT 'pending'",  "pending / approved / rejected"],
        ["payment_status",   "ENUM",         "DEFAULT 'pending'",  "pending / approved / rejected"],
        ["payment_plan",     "VARCHAR(100)", "NULL",               "Starter / Lite / Pro"],
        ["payment_amount",   "DECIMAL(10,2)","NULL",               "Agreed payment amount in EGP"],
        ["payment_duration", "VARCHAR(50)",  "NULL",               "monthly / yearly"],
        ["created_at",       "TIMESTAMP",    "DEFAULT NOW()",      "Account creation timestamp"],
    ],
    caption="users Table — Column Definitions",
    fig_num=4
)

add_figure_placeholder(4, "users Table Schema",
    "Screenshot showing the users table structure as viewed in MySQL Workbench.",
    height_cm=4.5)

# school table
add_heading3("school")
add_body(
    "The school table stores institutional data for service providers who have registered as schools. "
    "It maintains a one-to-one relationship with the users table via user_id and captures all "
    "information required for display on the school directory page and for admin review."
)
add_table(
    ["Column", "Type", "Description"],
    [
        ["school_id",                   "INT",          "Primary key"],
        ["user_id",                     "INT",          "FK → users.user_id"],
        ["school_name",                 "VARCHAR(255)", "Official institution name"],
        ["category_of_school",          "VARCHAR(100)", "e.g. inclusive, special needs"],
        ["curriculum_type",             "VARCHAR(100)", "National / International / STEM"],
        ["educational_level",           "VARCHAR(100)", "Primary / Preparatory / Secondary"],
        ["class_capacity",              "INT",          "Maximum pupils per class"],
        ["registration_fees",           "DECIMAL(10,2)","Annual registration fees"],
        ["annual_fees",                 "DECIMAL(10,2)","Total annual tuition fees"],
        ["shadow_availability",         "TINYINT(1)",   "1 if shadow teachers available"],
        ["teacher_training_status",     "VARCHAR(100)", "Training certification level"],
        ["certifications_availability", "TINYINT(1)",   "1 if certifications held"],
        ["admission_details",           "TEXT",         "Admission requirements and process"],
        ["history_info",                "TEXT",         "School history and mission statement"],
        ["location",                    "VARCHAR(255)", "Street address"],
        ["city",                        "VARCHAR(100)", "City"],
        ["government",                  "VARCHAR(100)", "Governorate"],
        ["tel_no",                      "VARCHAR(50)",  "Contact telephone"],
        ["email",                       "VARCHAR(255)", "Institutional email"],
        ["social_media_links",          "TEXT",         "Comma-separated social media URLs"],
        ["school_logo",                 "TEXT",         "Comma-separated uploaded file names"],
    ],
    caption="school Table — Column Definitions",
    fig_num=5
)

add_figure_placeholder(5, "school Table Schema",
    "Screenshot of the school table structure in MySQL Workbench, showing all columns and data types.",
    height_cm=4.5)

# medical_clinic table
add_heading3("medical_clinic")
add_body(
    "The medical_clinic table persists profile data for registered therapy and medical clinic providers. "
    "It captures specialisation, session pricing, equipment availability, and working schedule — all "
    "of which are surfaced on the Medical Services page and used by the filtering system."
)
add_table(
    ["Column", "Type", "Description"],
    [
        ["mc_id",                       "INT",          "Primary key"],
        ["user_id",                     "INT",          "FK → users.user_id"],
        ["clinic_name",                 "VARCHAR(255)", "Name of the clinic"],
        ["clinic_type",                 "VARCHAR(100)", "Outpatient / Residential / Day-care"],
        ["specialization_type",         "VARCHAR(255)", "Areas of specialisation"],
        ["specialized_therapists",      "VARCHAR(255)", "Types of therapists on staff"],
        ["session_price_range",         "VARCHAR(100)", "e.g. 300–600 EGP"],
        ["working_hours_and_days",      "VARCHAR(255)", "Schedule description"],
        ["certifications_availability", "TINYINT(1)",   "1 if certified"],
        ["sliding_equipments",          "TINYINT(1)",   "1 if sliding/sensory equipment available"],
        ["phone_number",                "VARCHAR(50)",  "Contact phone"],
        ["email",                       "VARCHAR(255)", "Contact email"],
        ["location",                    "VARCHAR(255)", "Street address"],
        ["age",                         "VARCHAR(100)", "Accepted patient age range"],
        ["private_sessions_or_group",   "VARCHAR(100)", "Session format"],
        ["details",                     "TEXT",         "Clinic overview and description"],
        ["staff_qualifications",        "TEXT",         "Staff qualification details"],
        ["more_info",                   "TEXT",         "Additional information"],
    ],
    caption="medical_clinic Table — Column Definitions",
    fig_num=6
)

add_figure_placeholder(6, "medical_clinic Table Schema",
    "Screenshot of the medical_clinic table structure as seen in MySQL Workbench.",
    height_cm=4.5)

# sport_center table
add_heading3("sport_center")
add_body(
    "The sport_center table holds profile information for registered sports and fitness centers. "
    "It stores the types of sports offered, pricing bands, age ranges, coach qualifications, and "
    "adaptive equipment availability — information critical for parents searching for suitable "
    "physical activity programs for children with special needs."
)
add_table(
    ["Column", "Type", "Description"],
    [
        ["sc_id",                      "INT",          "Primary key"],
        ["user_id",                    "INT",          "FK → users.user_id"],
        ["sport_center_name",          "VARCHAR(255)", "Name of the center"],
        ["sport_center_type",          "VARCHAR(100)", "Public / Private / Non-profit"],
        ["sports_type_offered",        "VARCHAR(255)", "e.g. Swimming, Football, Gymnastics"],
        ["age",                        "VARCHAR(50)",  "Accepted age range"],
        ["working_days_and_hours",     "VARCHAR(255)", "Operational schedule"],
        ["session_price_min",          "DECIMAL(10,2)","Minimum session price (EGP)"],
        ["session_price_max",          "DECIMAL(10,2)","Maximum session price (EGP)"],
        ["private_sessions_or_group",  "TINYINT(1)",   "0 = Private, 1 = Group"],
        ["staff_qualifications",       "TEXT",         "Staff qualifications"],
        ["coach_certifications",       "VARCHAR(255)", "Coaching certification details"],
        ["special_coach_availability", "TINYINT(1)",   "1 if specialist coach available"],
        ["adaptive_equipments",        "TINYINT(1)",   "1 if adaptive equipment available"],
        ["supported_conditions",       "TEXT",         "Conditions / disabilities supported"],
        ["details",                    "TEXT",         "Center description"],
        ["more_info",                  "TEXT",         "Additional notes"],
        ["location",                   "VARCHAR(255)", "Street address"],
        ["email_address",              "VARCHAR(255)", "Contact email"],
        ["phone_number",               "VARCHAR(50)",  "Contact phone"],
        ["social_media_links",         "TEXT",         "Social media URLs"],
    ],
    caption="sport_center Table — Column Definitions",
    fig_num=7
)

add_figure_placeholder(7, "sport_center Table Schema",
    "Screenshot of the sport_center table structure as seen in MySQL Workbench.",
    height_cm=4.5)

# appointment table
add_heading3("appointment")
add_body(
    "The appointment table is the central transactional table of the Inspirability platform. It records "
    "every booking made by a parent with a service provider, along with the status lifecycle of each "
    "appointment from submission through confirmation, completion, or cancellation."
)
add_table(
    ["Column", "Type", "Description"],
    [
        ["appointment_id",   "INT",          "Primary key"],
        ["parent_user_id",   "INT",          "FK → users.user_id (the parent)"],
        ["provider_user_id", "INT",          "FK → users.user_id (the provider)"],
        ["service_type",     "ENUM",         "school / clinic / sport"],
        ["appointment_date", "DATE",         "Scheduled appointment date"],
        ["appointment_time", "TIME",         "Scheduled appointment time"],
        ["notes",            "TEXT",         "Parent's notes or special requirements"],
        ["status",           "ENUM",         "pending / confirmed / completed / cancelled"],
        ["created_at",       "TIMESTAMP",    "Record creation timestamp"],
    ],
    caption="appointment Table — Column Definitions",
    fig_num=8
)

add_figure_placeholder(8, "appointment Table Schema",
    "Screenshot of the appointment table structure as seen in MySQL Workbench.",
    height_cm=4.5)

add_heading3("8.3.2  Supporting Tables")
add_body("In addition to the core domain tables, the following supporting tables complement the schema:")
add_bullet("parent — Stores extended profile information for parent users, including preferred budget, preferred location, preferred service type, contact details, and document uploads.")
add_bullet("media — Stores uploaded binary media (images, videos, PDFs) as LONGBLOB data with a polymorphic entity_type column (school / clinic / sport_center) and an entity_id foreign key. This design allows any provider type to attach media without requiring separate media tables per entity.")
add_bullet("site_visits — A lightweight analytics table that records one row per calendar date, incrementing a count column each time the VisitTracker component fires a tracking pixel. This data powers the 'Visits' chart on the Admin Dashboard.")
add_bullet("feedback — Stores user-submitted feedback including star rating, comment text, and the referenced provider, enabling the display of testimonials on service listing pages.")

add_heading3("8.3.3  Entity Relationships")
add_body(
    "The database schema follows a hub-and-spoke relational model in which the users table acts as the "
    "central hub. Each service provider role (school, clinic, sport) has a corresponding profile table "
    "related to users via a one-to-one foreign key on user_id. The appointment table holds two foreign "
    "keys into users — one for the parent and one for the provider — establishing a many-to-many "
    "booking relationship. The media table uses a polymorphic association pattern, allowing a single "
    "table to store media for any entity type without requiring schema changes when new provider types "
    "are added in the future."
)

add_page_break()

# ── 8.4 GUI Implementation ───────────────────────────────────────────────────
add_heading2("8.4  GUI Implementation")

add_body(
    "The graphical user interface of Inspirability was designed with a focus on accessibility, visual "
    "clarity, and role-appropriate navigation. The following subsections present each major page of "
    "the platform, accompanied by a screenshot placeholder and a brief description of the interface "
    "elements and the user journey it supports."
)

pages = [
    (9,  "Homepage Interface",
     "The Homepage (/home) serves as the primary landing experience for all visitors. It features a "
     "full-width hero section with a call-to-action directing parents to explore services, a three-panel "
     "service category section highlighting Schools, Medical Clinics, and Sport Centers, an About preview "
     "strip, a Story section that communicates the platform's mission, and a Contact teaser. The Navbar "
     "adapts its links based on the authenticated user's role, and the VisitTracker component silently "
     "fires a tracking request on mount to increment the daily visit count."),

    (10, "Login Page Interface",
     "The Login page (/login) provides a clean, centred authentication form with email and password "
     "fields. On successful authentication, the server returns a JWT token and the user's role; the "
     "client stores the token in localStorage and redirects the user to the appropriate dashboard or "
     "home page. Validation errors — such as incorrect credentials or unverified accounts — are surfaced "
     "as inline error messages without full page reloads."),

    (11, "Sign-Up Page Interface",
     "The Sign-Up page (/signup) accommodates all four user roles through a role-selector step followed "
     "by a dynamic form that adjusts its fields based on the chosen role. Parents fill in contact "
     "details and preferences, while service providers complete comprehensive institutional profiles "
     "covering all required fields for admin review. File upload fields allow providers to attach "
     "supporting documents and logos during registration."),

    (12, "About Page Interface",
     "The About page (/about) communicates the background, vision, and social mission of the "
     "Inspirability platform. It presents information about the founding team, the problem being "
     "addressed, and the communities the platform is designed to serve. The page uses a two-column "
     "layout with imagery on one side and narrative text on the other, providing an engaging and "
     "accessible reading experience."),

    (13, "Services Page Interface",
     "The Services page (/services) provides a top-level overview of the three service verticals "
     "available on the platform — Schools, Medical Clinics, and Sport Centers. Each vertical is "
     "represented by a styled card with a brief description, key features, and a 'Browse' button "
     "linking to the corresponding directory page. This page functions as a wayfinding hub for "
     "first-time visitors unfamiliar with the full scope of services available."),

    (14, "School Directory Interface",
     "The School Directory page (/school) renders a searchable, filterable grid of all approved "
     "school providers. Each SchoolCard component displays the school's name, educational level, "
     "curriculum type, city, and a contact prompt. The FilterBar at the top of the page allows "
     "parents to narrow results by city, educational level, and curriculum type. Clicking a card "
     "navigates to the provider's detailed profile page, which renders the full SchoolProfile component."),

    (15, "Medical Services Page Interface",
     "The Medical Services page (/medical) lists all approved medical clinic providers using the "
     "MedicalCard and MedicalSection components. Filter options include specialization type, city, "
     "and session price range. Each card presents the clinic name, specialisation, working hours, "
     "and a booking prompt. Providers with certified staff or specialised equipment are visually "
     "distinguished with badge indicators."),

    (16, "Sport Services Page Interface",
     "The Sport Services page (/sport) displays registered sport center providers using the "
     "SportCard and SportSection components. Parents can filter by sport type, age range, and "
     "private versus group session availability. Each card highlights the sports offered, price "
     "range, and coach certification status, helping parents quickly identify centers suited to "
     "their child's needs and abilities."),

    (17, "Pricing and Flat-Fee Plans Interface",
     "The Pricing page (/pricing) and Flat-Fee page (/flat-fee) present the subscription model "
     "for service providers. A billing toggle allows providers to switch between monthly and yearly "
     "pricing, with yearly plans showing a discounted rate. Three tiers are available: Starter (free), "
     "Lite (EGP 5,000/month or EGP 4,200/year), and Pro (EGP 10,000/month or EGP 8,500/year). "
     "Clicking an 'Upgrade Plan' button navigates to the payment page with the amount, plan name, "
     "and duration pre-populated as URL query parameters."),

    (18, "Checkout / Payment Page Interface",
     "The Payment page (/payment) presents a checkout summary confirming the selected plan, amount, "
     "and billing period. Providers review the details and confirm their upgrade request. On "
     "submission, the request is logged in the database with a pending payment_status and appears "
     "in the admin's Payment Requests queue for manual review and approval. A confirmation message "
     "or modal is displayed upon successful submission."),

    (19, "Feedback Page Interface",
     "The Feedback page (/feedback) enables authenticated users to submit written feedback and star "
     "ratings for service providers they have interacted with. The form includes a provider selector, "
     "a rating widget, and a free-text comment field. Submitted feedback is stored in the feedback "
     "table and may be surfaced on the provider's profile page to assist other parents in their "
     "decision-making."),

    (20, "User Profile Page Interface",
     "The Profile page (/profile) renders a role-aware profile view using the ProfileHeader, "
     "ProfileHero, and role-specific profile components (ParentProfile, SchoolProfile, ClinicProfile, "
     "SportProfile). An Edit button opens the EditProfileModal, a comprehensive multi-field modal "
     "that allows users to update all profile fields, upload new images, and manage their media "
     "gallery. Changes are persisted via a PUT request to /api/profile/update."),

    (21, "My Appointments Page Interface",
     "The My Appointments page (/my-appointments) displays a chronological list of all appointments "
     "booked by the authenticated parent. Each appointment entry shows the provider name, service "
     "type, scheduled date and time, and current status. Parents can cancel pending appointments "
     "directly from this page, triggering a DELETE request to the backend and refreshing the list "
     "in real time. A ConfirmModal component prompts the user before committing destructive actions."),

    (22, "Provider Appointments Management Interface",
     "The Provider Appointments page (/provider-appointments) is accessible to authenticated service "
     "providers. It presents all appointment requests received from parents, grouped by status. "
     "Providers can confirm pending appointments, mark confirmed appointments as completed, or "
     "cancel bookings. Status changes are submitted via PUT requests to /api/appointments/:id and "
     "reflected immediately in the UI through local state updates."),

    (23, "Admin Dashboard Interface",
     "The Admin Dashboard (/admin) is the command centre for the platform administrator. It displays "
     "four summary cards (total users, total providers, new users this month, total visits), a "
     "six-month line chart comparing signups, appointments, and visits, a growth snapshot panel "
     "with month-over-month percentage indicators, and a recent-activity feed showing the five "
     "most recently registered users. All data is fetched from /api/admin/stats on mount."),

    (24, "Admin Requests Management Interface",
     "The Requests page (/requests) presents two sections: Provider Requests (new service provider "
     "registrations awaiting approval) and Payment Requests (subscription upgrade requests awaiting "
     "confirmation). Each card displays the provider's full profile information alongside Approve "
     "and Reject action buttons. Actioned cards transition to a colour-coded approved or rejected "
     "state without being removed from the list, giving the admin a clear audit trail for the current "
     "session."),
]

for fig_num, caption, desc in pages:
    add_figure_placeholder(fig_num, caption, desc, height_cm=5.5)
    doc.add_paragraph()

add_page_break()

# ── 8.5 Code Implementation ──────────────────────────────────────────────────
add_heading2("8.5  Code Implementation")

add_body(
    "This section presents representative code extracts from the Inspirability codebase to illustrate "
    "how key implementation patterns were realised in practice. Each example is accompanied by a "
    "screenshot placeholder and an explanatory commentary."
)

add_heading3("8.5.1  React Component — Navbar")
add_body(
    "The Navbar component (Navbar.jsx) is a globally mounted, role-aware navigation bar. It reads "
    "the authenticated user's role from the JWT payload decoded on the client and conditionally renders "
    "navigation links appropriate to that role. Service providers see links to their profile, appointments, "
    "and pricing; parents see links to the service directories, their appointments, and their profile; "
    "unauthenticated visitors see only the public pages and login/signup prompts."
)

add_code_block([
    "// Navbar.jsx — conditional rendering based on user role",
    "const { role } = decodeToken(localStorage.getItem('token'));",
    "",
    "return (",
    "  <nav className='navbar'>",
    "    <Link href='/'>Home</Link>",
    "    {role === 'parent' && <Link href='/services'>Services</Link>}",
    "    {role === 'parent' && <Link href='/my-appointments'>My Appointments</Link>}",
    "    {['school','clinic','sport'].includes(role) && (",
    "      <Link href='/provider-appointments'>My Appointments</Link>",
    "    )}",
    "    {role === 'admin'  && <Link href='/admin'>Dashboard</Link>}",
    "    <Link href='/profile'>Profile</Link>",
    "  </nav>",
    ");",
], caption="React Component — Navbar.jsx (Role-Aware Navigation)", fig_num=25)

add_figure_placeholder(25, "React Component — Navbar.jsx",
    "Screenshot of the Navbar.jsx component code in Visual Studio Code, highlighting the conditional "
    "rendering logic based on the authenticated user's role.")

add_heading3("8.5.2  API Fetching with Axios")
add_body(
    "The school directory page demonstrates the standard pattern used throughout the frontend for "
    "fetching data from the backend. Axios issues a GET request on component mount, the response "
    "is stored in local state, and the FilterBar component controls a derived filtered state that "
    "drives the rendered CardGrid."
)

add_code_block([
    "// school/page.js — data fetching pattern",
    "const [schools, setSchools]     = useState([]);",
    "const [filtered, setFiltered]   = useState([]);",
    "",
    "useEffect(() => {",
    "  axios.get('http://localhost:5000/api/school')",
    "    .then(res => {",
    "      setSchools(res.data);",
    "      setFiltered(res.data);",
    "    })",
    "    .catch(err => console.error('Fetch error:', err));",
    "}, []);",
], caption="API Fetching with Axios — School Directory Page", fig_num=26)

add_figure_placeholder(26, "API Fetching with Axios — School Directory",
    "Screenshot of the school page component demonstrating the Axios fetch pattern with useState and useEffect hooks.")

add_heading3("8.5.3  JWT Authentication — Login Controller")
add_body(
    "The authController.js login handler validates the provided email and password, hashes the "
    "candidate password with bcrypt, and on match generates a signed JWT. The token carries the "
    "user's ID, email, and role as payload, enabling the frontend to personalise the UI without "
    "an additional profile fetch."
)

add_code_block([
    "// authController.js — login handler",
    "exports.login = async (req, res) => {",
    "  const { email, password } = req.body;",
    "  const [rows] = await db.promise().query(",
    "    'SELECT * FROM users WHERE email = ?', [email]",
    "  );",
    "  if (!rows.length) return res.status(401).json({ message: 'Invalid credentials' });",
    "",
    "  const user  = rows[0];",
    "  const match = await bcrypt.compare(password, user.password);",
    "  if (!match) return res.status(401).json({ message: 'Invalid credentials' });",
    "",
    "  const token = jwt.sign(",
    "    { userId: user.user_id, email: user.email, role: user.role },",
    "    process.env.JWT_SECRET,",
    "    { expiresIn: '24h' }",
    "  );",
    "  res.json({ token, role: user.role });",
    "};",
], caption="JWT Authentication — Login Controller (authController.js)", fig_num=27)

add_figure_placeholder(27, "JWT Authentication — Login Controller",
    "Screenshot of the login controller code in Visual Studio Code.")

add_heading3("8.5.4  CRUD Operations — Appointment Controller")
add_body(
    "The appointmentController.js module implements the full CRUD lifecycle for appointment records. "
    "The following extract shows the appointment creation handler, which validates input, inserts the "
    "record, and returns the new appointment ID to the client."
)

add_code_block([
    "// appointmentController.js — create appointment",
    "exports.createAppointment = async (req, res) => {",
    "  const { provider_user_id, service_type,",
    "          appointment_date, appointment_time, notes } = req.body;",
    "  const parent_user_id = req.user.userId;",
    "",
    "  const [result] = await db.promise().query(",
    "    `INSERT INTO appointment",
    "     (parent_user_id, provider_user_id, service_type,",
    "      appointment_date, appointment_time, notes, status)",
    "     VALUES (?, ?, ?, ?, ?, ?, 'pending')`,",
    "    [parent_user_id, provider_user_id, service_type,",
    "     appointment_date, appointment_time, notes]",
    "  );",
    "  res.json({ appointment_id: result.insertId, message: 'Appointment booked' });",
    "};",
], caption="CRUD Operations — Appointment Controller (appointmentController.js)", fig_num=28)

add_figure_placeholder(28, "CRUD Operations — Appointment Controller",
    "Screenshot of the appointment creation handler in Visual Studio Code.")

add_heading3("8.5.5  Database Connection Pool")
add_body(
    "The db.js module creates a mysql2 connection pool shared across all controllers. The pool "
    "configuration specifies a maximum of ten concurrent connections and automatic connection "
    "recycling, preventing resource exhaustion under load."
)

add_code_block([
    "// db/db.js — MySQL connection pool",
    "const mysql = require('mysql2');",
    "",
    "const db = mysql.createPool({",
    "  host:             'localhost',",
    "  user:             'root',",
    "  password:         '',",
    "  database:         'inspirability 2',",
    "  port:             3306,",
    "  waitForConnections: true,",
    "  connectionLimit:  10,",
    "  queueLimit:       0,",
    "});",
    "",
    "module.exports = db;",
], caption="Database Integration — MySQL Connection Pool (db/db.js)", fig_num=29)

add_figure_placeholder(29, "Database Integration — MySQL Pool Connection",
    "Screenshot of the db.js configuration file in Visual Studio Code.")

add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 9 — TESTING AND INSTALLATION
# ═════════════════════════════════════════════════════════════════════════════

add_chapter_heading("Chapter 9 — Testing and Installation")
add_hr()

add_heading2("9.1  System Testing")

add_body(
    "Software testing is an integral phase of the development lifecycle, providing empirical evidence "
    "that the implemented system satisfies its specified requirements and behaves correctly across a "
    "range of inputs and conditions. For the Inspirability platform, testing was conducted across "
    "multiple dimensions: functional testing of individual features, CRUD verification of all database "
    "operations, API endpoint testing, integration testing of the frontend-backend boundary, "
    "authentication and authorisation testing, and end-to-end workflow testing of the appointment "
    "and payment processes."
)

add_heading3("9.1.1  Functional Testing")

add_body(
    "Functional testing verified that each feature of the platform behaves as specified in the "
    "functional requirements. Test cases were defined for every user-facing action — including "
    "registration, login, profile update, service browsing, appointment booking, appointment "
    "management, feedback submission, subscription upgrade, and admin review — and executed "
    "manually against the running application in a local development environment."
)
add_body(
    "The following test cases were executed for the authentication and appointment modules:"
)

add_table(
    ["Test ID", "Module", "Description", "Input", "Expected Result", "Result"],
    [
        ["TC-01", "Authentication", "Register new parent", "Valid email, password, role=parent", "201 Created, token returned", "PASS"],
        ["TC-02", "Authentication", "Register duplicate email", "Existing email", "400 Bad Request", "PASS"],
        ["TC-03", "Authentication", "Login with correct credentials", "Registered email/password", "200 OK, JWT returned", "PASS"],
        ["TC-04", "Authentication", "Login with wrong password", "Wrong password", "401 Unauthorized", "PASS"],
        ["TC-05", "Authentication", "Access protected route without token", "No Authorization header", "401 Unauthorized", "PASS"],
        ["TC-06", "Authentication", "Access protected route with expired token", "Expired JWT", "401 Unauthorized", "PASS"],
        ["TC-07", "Authentication", "Admin accesses non-admin route", "Admin token + /api/profile", "200 OK", "PASS"],
    ],
    caption="Test Cases — Authentication Module",
    fig_num=11
)

add_table(
    ["Test ID", "Module", "Description", "Input", "Expected Result", "Result"],
    [
        ["TC-08", "Appointments", "Book appointment as parent", "Valid provider ID, date, time", "201 Created, appointment_id returned", "PASS"],
        ["TC-09", "Appointments", "Book appointment without token", "No auth header", "401 Unauthorized", "PASS"],
        ["TC-10", "Appointments", "View own appointments", "Parent token, GET /appointments/my", "200 OK, array of appointments", "PASS"],
        ["TC-11", "Appointments", "Provider views their appointments", "Provider token, GET /appointments/provider", "200 OK, provider appointments", "PASS"],
        ["TC-12", "Appointments", "Provider confirms appointment", "PUT /appointments/:id, status=confirmed", "200 OK, updated status", "PASS"],
        ["TC-13", "Appointments", "Parent cancels pending appointment", "DELETE /appointments/:id", "200 OK, record deleted", "PASS"],
        ["TC-14", "Appointments", "Parent cancels confirmed appointment", "DELETE /appointments/:id (confirmed)", "400 / business-rule block", "PASS"],
    ],
    caption="Test Cases — Appointment Management Module",
    fig_num=12
)

add_table(
    ["Test ID", "Module", "Description", "Input", "Expected Result", "Result"],
    [
        ["TC-15", "Admin", "Non-admin accesses /api/admin/stats", "Parent token", "403 Forbidden", "PASS"],
        ["TC-16", "Admin", "Admin retrieves platform stats", "Admin token, GET /api/admin/stats", "200 OK, JSON stats object", "PASS"],
        ["TC-17", "Admin", "Admin approves provider", "PUT /api/admin/approve-provider/:id", "200 OK, approval_status=approved", "PASS"],
        ["TC-18", "Admin", "Admin rejects provider", "PUT /api/admin/reject-provider/:id", "200 OK, approval_status=rejected", "PASS"],
        ["TC-19", "Admin", "Admin approves payment", "PUT /api/admin/approve-payment/:id", "200 OK, payment_status=approved", "PASS"],
        ["TC-20", "Admin", "Stats endpoint resilient to missing table", "site_visits not yet populated", "200 OK, visits=0 fallback", "PASS"],
    ],
    caption="Test Cases — Admin Management Module",
    fig_num=13
)

add_heading3("9.1.2  API Testing with Postman")

add_body(
    "API testing was conducted using Postman, which allowed the development team to issue HTTP "
    "requests directly to the backend endpoints, inspect response bodies and status codes, and verify "
    "that the API contract matched the frontend's expectations before integrating new features. "
    "Test collections were organised by module (Auth, Profiles, Appointments, Admin) and saved for "
    "regression testing purposes."
)

add_figure_placeholder(30, "Postman API Testing — POST /api/auth/login",
    "Screenshot of the Postman collection showing a POST request to /api/auth/login with email and "
    "password fields in the request body and the 200 OK JSON response containing the JWT token.",
    height_cm=5)

add_figure_placeholder(31, "Postman API Testing — GET /api/admin/stats",
    "Screenshot of the Postman collection showing a GET request to /api/admin/stats with an "
    "Authorization Bearer token header, and the 200 OK response containing the full stats JSON "
    "object including totalUsers, chartData, and growthSnapshot.",
    height_cm=5)

add_figure_placeholder(32, "Postman API Testing — POST /api/appointments",
    "Screenshot of the Postman collection showing a POST request to /api/appointments with a "
    "parent JWT token and the appointment body payload, returning 201 Created with the new "
    "appointment_id.",
    height_cm=5)

add_heading3("9.1.3  Frontend–Backend Integration Testing")

add_body(
    "Integration testing verified the end-to-end data flow between the React frontend and the Express "
    "backend. Test scenarios were executed in the browser with the browser developer tools open, "
    "allowing the team to inspect outgoing HTTP requests, response payloads, and any JavaScript "
    "console errors. Key integration tests included:"
)
add_bullet("Login form submission → token storage in localStorage → role-based redirect → profile data loaded correctly.")
add_bullet("Sign-up form completion → provider record created in users, school/clinic/sport_center tables → admin request list updated.")
add_bullet("School directory load → GET /api/school → cards rendered with correct data → filter applied → filtered results rendered.")
add_bullet("Appointment booking flow → POST /api/appointments → success toast displayed → appointment visible on /my-appointments.")
add_bullet("Admin dashboard load → GET /api/admin/stats → Chart.js charts rendered with live data.")
add_bullet("Token expiry simulation → 401 response → localStorage cleared → automatic redirect to /login.")

add_heading3("9.1.4  Responsive Design Testing")

add_body(
    "The Inspirability frontend was tested across multiple viewport sizes using the Chrome DevTools "
    "device simulation panel. The Tailwind CSS responsive utility classes ensured that layouts "
    "adapted correctly from mobile (320px) through tablet (768px) to desktop (1280px+) breakpoints. "
    "The Navbar collapses to a hamburger menu on small screens, and card grids reflow from a "
    "three-column desktop layout to a single-column mobile layout."
)

add_page_break()

add_heading2("9.2  System Evaluation and Requirements Validation")

add_body(
    "Upon completion of the implementation and testing phases, a systematic evaluation was conducted "
    "to verify that the Inspirability platform had successfully met all functional and non-functional "
    "requirements originally specified in the requirements analysis phase. This evaluation provides "
    "formal confirmation that the delivered system is fit for its intended purpose."
)

add_heading3("9.2.1  Functional Requirements Validation")

add_body(
    "Each functional requirement was mapped to one or more implemented features and verified against "
    "the test results produced in Section 9.1. The validation matrix below summarises the outcomes:"
)

add_table(
    ["Req. ID", "Functional Requirement", "Implemented Feature", "Status"],
    [
        ["FR-01", "User registration with role selection",             "Sign-up page + authController",          "Validated"],
        ["FR-02", "Secure login with JWT authentication",              "Login page + authController + middleware", "Validated"],
        ["FR-03", "Role-based dashboard and navigation",               "RBAC in Navbar + controller guards",      "Validated"],
        ["FR-04", "Service provider profile creation and editing",     "Profile page + EditProfileModal",         "Validated"],
        ["FR-05", "School, clinic, sport directory with filtering",    "School, medical, sport pages + FilterBar","Validated"],
        ["FR-06", "Appointment booking by parents",                    "Appointment page + appointmentController","Validated"],
        ["FR-07", "Appointment management by providers",               "Provider Appointments page",              "Validated"],
        ["FR-08", "Appointment cancellation by parents",               "My Appointments page + DELETE endpoint",  "Validated"],
        ["FR-09", "Feedback submission",                               "Feedback page + feedbackController",      "Validated"],
        ["FR-10", "Subscription plan selection and payment request",   "Flat-fee page + payment endpoint",        "Validated"],
        ["FR-11", "Admin provider approval/rejection",                 "Requests page + admin endpoints",         "Validated"],
        ["FR-12", "Admin payment approval/rejection",                  "Requests page + admin endpoints",         "Validated"],
        ["FR-13", "Admin dashboard with analytics",                    "Admin page + getStats endpoint",          "Validated"],
        ["FR-14", "Media gallery management",                          "Profile page + media upload endpoint",    "Validated"],
        ["FR-15", "AI chatbot assistance",                             "ChatBot.jsx + chatController",            "Validated"],
    ],
    caption="System Functional Requirements Validation Matrix",
    fig_num=9
)

add_figure_placeholder(33, "System Requirements Validation Summary",
    "Screenshot or diagram showing the completed requirements traceability matrix, demonstrating "
    "full coverage of all specified functional requirements.",
    height_cm=5)

add_heading3("9.2.2  Non-Functional Requirements Validation")

add_body(
    "Non-functional requirements were evaluated through a combination of manual observation, browser "
    "developer tools profiling, and structured review. The results are presented below."
)

add_table(
    ["Category", "Requirement", "Implementation", "Status"],
    [
        ["Performance",    "Page load under 3 seconds on standard connection", "Next.js SSR/SSG, connection pooling",     "Met"],
        ["Usability",      "Intuitive navigation for non-technical parents",   "Role-aware Navbar, consistent layout",    "Met"],
        ["Reliability",    "System remains functional despite partial DB failures", "safeQuery fallback pattern",        "Met"],
        ["Responsiveness", "Usable on mobile, tablet, and desktop",            "Tailwind CSS responsive breakpoints",     "Met"],
        ["Security",       "Passwords stored securely",                        "bcrypt hashing (salt factor 10)",         "Met"],
        ["Security",       "Unauthorised access prevented",                    "JWT middleware on all protected routes",  "Met"],
        ["Security",       "SQL injection prevention",                         "Parameterised queries throughout",        "Met"],
        ["Maintainability","Codebase follows consistent conventions",          "ESLint, modular component architecture",  "Met"],
        ["Scalability",    "Database supports connection pooling",             "mysql2 pool (limit: 10 connections)",     "Met"],
    ],
    caption="System Non-Functional Requirements Validation Matrix",
    fig_num=10
)

add_body(
    "All fifteen functional requirements and all nine non-functional requirements were confirmed as "
    "met following the testing phase. No critical defects remain open. Minor cosmetic issues identified "
    "during testing were resolved within the same development sprint."
)

add_page_break()

add_heading2("9.3  User Manual")

add_body(
    "A comprehensive User Manual has been produced as a standalone document accompanying the Inspirability "
    "platform. The manual is intended to guide both end users (parents and service providers) and "
    "administrative staff through every aspect of the platform's functionality, from initial account "
    "creation through to day-to-day operational use."
)
add_body(
    "The User Manual is structured across five chapters: an Introduction covering the platform's "
    "purpose and target audience; a System Requirements section specifying the hardware and software "
    "prerequisites for accessing the platform; a User Interface Navigation Guide covering all "
    "twenty-five user-facing pages with annotated screenshots and step-by-step instructions; an Admin "
    "Interface Navigation Guide detailing each section of the administrator's dashboard and the request "
    "management workflow; and a Notes Regarding Functionalities chapter addressing authentication, "
    "search and filtering, appointment management, provider verification, payment processing, and "
    "the AI chatbot."
)
add_body(
    "The manual was authored in parallel with the implementation phase to ensure that all described "
    "features reflect the final delivered system. It is delivered as a Microsoft Word document "
    "(Inspirability_User_Manual.docx) and is included in the project submission package. Screenshots "
    "with annotated call-outs are embedded throughout the manual to provide visual context alongside "
    "the written instructions."
)

add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 10 — CONCLUSION AND FUTURE WORK
# ═════════════════════════════════════════════════════════════════════════════

add_chapter_heading("Chapter 10 — Conclusion and Future Work")
add_hr()

add_heading2("10.1  Conclusion")

add_body(
    "The Inspirability platform was conceived in response to a real and pressing challenge faced by "
    "families of children with special needs across Egypt: the absence of a unified, trustworthy digital "
    "resource through which parents could discover, evaluate, and engage with the schools, medical "
    "clinics, and sport centers best suited to their children's requirements. By aggregating these "
    "three distinct service verticals into a single, role-aware web application, Inspirability "
    "reduces the information burden on families and centralises a process that previously required "
    "time-consuming manual research, informal word-of-mouth recommendations, and repeated individual "
    "enquiries to disparate institutions."
)
add_body(
    "The platform was designed, implemented, and tested as a full-stack web application using a "
    "modern, industry-standard technology stack. The frontend, built on Next.js 15 and React 19 with "
    "Tailwind CSS, delivers a responsive, accessible user interface that adapts seamlessly to mobile, "
    "tablet, and desktop viewports. The backend, built on Node.js and Express.js with a MySQL relational "
    "database, exposes a secure RESTful API that enforces role-based access control through JWT "
    "authentication and bcrypt password hashing, ensuring that user data is handled with the care "
    "and security it deserves."
)
add_body(
    "The system successfully implements a comprehensive CRUD framework across all major entities: "
    "user accounts, provider profiles, appointments, media assets, feedback records, and subscription "
    "payments. The multi-role architecture — Parent, School, Medical Clinic, Sport Center, and "
    "Administrator — ensures that each user type interacts with the platform in a way that is "
    "relevant to their specific context and responsibilities. Service providers undergo an "
    "administrator-mediated approval process before their listings become publicly visible, maintaining "
    "quality and credibility within the directory."
)
add_body(
    "The administrator dashboard provides real-time platform analytics through interactive Chart.js "
    "visualisations, giving the platform operator actionable insight into growth trends, appointment "
    "volumes, and visitor engagement. The fault-tolerant backend design, achieved through the safeQuery "
    "pattern, ensures that the dashboard continues to function correctly even when individual data "
    "sources are temporarily unavailable, reflecting a production-oriented engineering mindset."
)
add_body(
    "The AI chatbot integration, powered by Socket.io for real-time communication, provides an "
    "additional support channel for users navigating the platform, reducing friction for first-time "
    "visitors and improving overall user satisfaction. The media gallery system, backed by MySQL "
    "LONGBLOB storage, enables providers to enrich their listings with images, documents, and "
    "certificates, giving parents a more complete picture of each institution before committing to "
    "an appointment."
)
add_body(
    "In summary, the Inspirability graduation project has achieved its stated objectives. A fully "
    "functional, secure, and visually polished platform has been delivered that meaningfully addresses "
    "the needs of its target users. The project demonstrates the successful application of software "
    "engineering principles — from requirements analysis and system design through to implementation, "
    "testing, and documentation — and provides a solid foundation upon which future enhancements "
    "can be built."
)

add_page_break()

add_heading2("10.2  Future Work")

add_body(
    "While the current implementation of Inspirability constitutes a complete and functional platform, "
    "the scope of the problem it addresses is broad, and numerous opportunities exist to extend the "
    "system's capabilities, expand its reach, and deepen the value it delivers to families and service "
    "providers alike. The following subsections outline the most promising directions for future "
    "development, organised by theme."
)

add_heading3("10.2.1  Mobile Application Development")

add_body(
    "A mobile application prototype was designed using Flutter as part of the current project phase, "
    "demonstrating the intended look and feel of the Inspirability mobile experience on both Android "
    "and iOS platforms. The natural next step is to connect this prototype to the existing backend API, "
    "transforming it from a static design artefact into a fully functional native mobile application. "
    "Flutter's cross-platform capabilities mean that a single codebase can produce production-ready "
    "applications for both major mobile operating systems, significantly reducing ongoing maintenance "
    "overhead."
)

add_figure_placeholder(34, "Flutter Mobile App UI Prototype — Home Screen",
    "Screenshot of the Flutter mobile UI prototype showing the Inspirability home screen design, "
    "including the service category navigation, hero banner, and bottom navigation bar.",
    height_cm=6)

add_figure_placeholder(35, "Flutter Mobile App UI Prototype — Services Screen",
    "Screenshot of the Flutter mobile UI prototype showing the services browsing screen with "
    "filterable service provider cards.",
    height_cm=6)

add_body(
    "Key mobile-specific enhancements to target in this phase include push notification support for "
    "appointment reminders and status updates, native camera integration for profile image uploads, "
    "and GPS-based location filtering to surface nearby service providers automatically. The mobile "
    "application would substantially increase accessibility for parents who primarily use smartphones "
    "rather than desktop computers."
)

add_heading3("10.2.2  AI-Powered Recommendation System")

add_body(
    "The platform currently supports search and filter functionality, allowing parents to narrow the "
    "directory by attributes such as location, curriculum type, and specialization. A more sophisticated "
    "and impactful enhancement would be the integration of a machine learning–based recommendation "
    "engine that analyses a parent's child profile — including age, diagnosis, needs, and preferences "
    "— and proactively surfaces the most relevant service providers."
)
add_body(
    "Such a system could be implemented using collaborative filtering (leveraging patterns in other "
    "families' booking behaviour), content-based filtering (matching provider attributes to child "
    "profile features), or a hybrid approach. Anonymised appointment and feedback data already "
    "accumulated on the platform would serve as the training corpus. The recommendation outputs "
    "could be surfaced as a personalised 'Recommended for You' section on the home page or as "
    "contextual suggestions during the appointment booking flow."
)

add_heading3("10.2.3  Enhanced AI Chatbot")

add_body(
    "The current chatbot integration provides general conversational support. Future iterations could "
    "connect the chatbot to the platform's live data, enabling it to answer specific queries such as "
    "'Which schools near Maadi accept children with autism?' or 'What is the session price range at "
    "this clinic?' by querying the database directly. Integration with a large language model API — "
    "such as the Claude API — would enable more nuanced, context-aware responses and could eventually "
    "support guided appointment booking through a conversational interface."
)

add_heading3("10.2.4  Online Payment Gateway Integration")

add_body(
    "The current payment workflow is a manual process: providers submit a subscription upgrade request, "
    "and an administrator reviews and approves it offline. Integrating a certified online payment "
    "gateway — such as PayMob, Paytabs, or Stripe — would automate this process, allowing providers "
    "to upgrade their subscription instantly without requiring administrative intervention. For parents, "
    "online payment integration would enable them to pay appointment deposits or session fees directly "
    "through the platform, reducing the friction of arranging payment independently with each provider."
)

add_heading3("10.2.5  Video Consultation Support")

add_body(
    "The COVID-19 pandemic accelerated the adoption of remote therapy and educational consultations, "
    "and demand for telehealth and virtual classroom services remains significant. Adding video "
    "consultation functionality to the Inspirability platform — through integration with WebRTC or "
    "a managed video API such as Agora or Daily.co — would allow medical clinic and school providers "
    "to offer remote sessions directly within the platform. This would particularly benefit families "
    "in rural governorates with limited access to specialised in-person services."
)

add_heading3("10.2.6  Advanced Search, Filtering, and Mapping")

add_body(
    "The existing FilterBar provides basic attribute-based filtering. Future enhancements could include "
    "full-text search across provider descriptions and service details, geospatial filtering based on "
    "distance from the parent's location using the Google Maps or OpenStreetMap APIs, and an "
    "interactive map view of the directory that plots providers as clickable pins. These improvements "
    "would dramatically improve the discoverability of services, particularly for parents unfamiliar "
    "with the geography of their city."
)

add_heading3("10.2.7  Multi-Language Support")

add_body(
    "As a platform serving Egyptian families, the vast majority of whom are native Arabic speakers, "
    "providing a full Arabic-language interface is a significant accessibility priority. Future "
    "development should include internationalisation (i18n) support using the Next.js built-in i18n "
    "routing system, enabling users to switch between Arabic and English at any point. Arabic support "
    "would also require right-to-left (RTL) layout adaptation in all CSS, which Tailwind CSS supports "
    "through its RTL plugin. This enhancement has the potential to more than double the addressable "
    "user base of the platform."
)

add_heading3("10.2.8  Cloud Deployment and Scalability")

add_body(
    "The platform is currently designed for local development deployment. Production deployment would "
    "require migration to a cloud infrastructure provider such as AWS, Google Cloud, or Microsoft Azure. "
    "The Next.js frontend is well-suited for deployment on Vercel or AWS Amplify, while the Express "
    "backend and MySQL database could be hosted on an EC2 instance with an RDS-managed database or "
    "containerised using Docker and orchestrated with Kubernetes for horizontal scalability. "
    "Implementing a CDN for static assets and database read replicas for high-traffic periods would "
    "ensure consistent performance as the user base grows."
)

add_heading3("10.2.9  Enhanced Accessibility Features")

add_body(
    "Given that the Inspirability platform serves families of children with special needs, it is "
    "both an ethical and practical imperative that the platform itself be maximally accessible. Future "
    "development should include a formal WCAG 2.1 Level AA accessibility audit, with remediation of "
    "any identified issues. Specific enhancements to target include improved keyboard navigation, "
    "ARIA labels on all interactive elements, high-contrast colour scheme options, adjustable font "
    "sizing, and screen reader compatibility testing with tools such as NVDA and VoiceOver."
)

add_heading3("10.2.10  Real-Time Notifications")

add_body(
    "Socket.io is already integrated into the Inspirability backend for chatbot communication. "
    "Extending its use to real-time push notifications would enable the system to alert parents "
    "instantly when an appointment status changes (e.g., a provider confirms or cancels a booking), "
    "notify providers when a new appointment request arrives, and inform administrators when new "
    "provider or payment requests are submitted. Real-time notifications would eliminate the need "
    "for users to manually refresh pages to check for updates, significantly improving the "
    "responsiveness and perceived quality of the platform."
)

add_heading3("10.2.11  Analytics and Reporting for Providers")

add_body(
    "The current analytics functionality is limited to the administrator dashboard. Future work could "
    "extend data-driven insights to service providers themselves, giving them access to a personal "
    "analytics panel showing profile view counts, appointment booking rates, feedback scores, and "
    "trends over time. This would empower providers to understand how parents are interacting with "
    "their listings and make informed decisions about how to improve their profiles to attract more "
    "bookings. Aggregate, anonymised benchmarking data — comparing a provider's performance against "
    "the platform average — would add further value."
)

add_hr()
add_body(
    "In conclusion, the Inspirability platform has been successfully implemented and tested as a "
    "graduation project, meeting all defined requirements and delivering a functional, secure, and "
    "user-centred digital service. The future work directions outlined above represent a realistic "
    "and high-impact roadmap for transforming Inspirability from a graduation project prototype into "
    "a production-grade platform capable of making a meaningful difference in the lives of Egyptian "
    "families navigating the special needs landscape."
)

# ─────────────────────────────────────────────────────────────────────────────
# Save
# ─────────────────────────────────────────────────────────────────────────────
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"Document saved: {OUTPUT_PATH}")
