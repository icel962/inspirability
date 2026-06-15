"""
generate_features_guide.py
Generates a clean Feature Reference Guide for the Inspirability Graduation Project.
Covers the 7 core features with frontend, backend, database, and screenshot guidance.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Colors ──────────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x16, 0x22, 0x7E)
TEAL      = RGBColor(0x00, 0xA6, 0xC8)
DARK      = RGBColor(0x1A, 0x1A, 0x2E)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GREEN_TXT = RGBColor(0x15, 0x57, 0x24)
GRAY_TXT  = RGBColor(0x44, 0x44, 0x44)

OUTPUT_PATH = os.path.join(
    os.path.dirname(__file__),
    "my-app",
    "Inspirability_Features_Guide.docx",
)

doc = Document()

# ── Page setup ───────────────────────────────────────────────────────────────
for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.5)

doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(12)

# ── Helpers ──────────────────────────────────────────────────────────────────

def cell_shd(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    s    = OxmlElement("w:shd")
    s.set(qn("w:val"),   "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"),  hex_color)
    tcPr.append(s)


def para_shd(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    s   = OxmlElement("w:shd")
    s.set(qn("w:val"),   "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"),  hex_color)
    pPr.append(s)


def hr(color="16227E"):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr= OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "8")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def title_block(text, subtitle=""):
    """Big navy title block at top of document."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    para_shd(p, "16227E")
    r = p.add_run("\n  " + text + "\n")
    r.font.name  = "Times New Roman"
    r.font.size  = Pt(22)
    r.font.bold  = True
    r.font.color.rgb = WHITE
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(16)
        para_shd(p2, "1A2A6C")
        r2 = p2.add_run("  " + subtitle + "  ")
        r2.font.name   = "Times New Roman"
        r2.font.size   = Pt(12)
        r2.font.italic = True
        r2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xFF)


def feature_header(number, name, emoji=""):
    """Teal banner for each feature section."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(2)
    para_shd(p, "00A6C8")
    r = p.add_run(f"   Feature {number}   {emoji}  {name}")
    r.font.name  = "Times New Roman"
    r.font.size  = Pt(14)
    r.font.bold  = True
    r.font.color.rgb = WHITE


def section_label(text):
    """Small navy label (e.g. DESCRIPTION, FRONTEND REFERENCES)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    para_shd(p, "EEF0FF")
    r = p.add_run(f"  {text}")
    r.font.name  = "Times New Roman"
    r.font.size  = Pt(10)
    r.font.bold  = True
    r.font.color.rgb = NAVY


def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.5)
    r = p.add_run(text)
    r.font.name  = "Times New Roman"
    r.font.size  = Pt(11)
    r.font.color.rgb = DARK


def ref_table(frontend_rows, backend_rows, db_rows):
    """Three-column reference table: Frontend | Backend | Database."""
    max_rows = max(len(frontend_rows), len(backend_rows), len(db_rows))
    # pad
    frontend_rows = frontend_rows + [""] * (max_rows - len(frontend_rows))
    backend_rows  = backend_rows  + [""] * (max_rows - len(backend_rows))
    db_rows       = db_rows       + [""] * (max_rows - len(db_rows))

    t = doc.add_table(rows=1 + max_rows, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for row in t.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(5.5)

    headers = ["Frontend Files / Pages", "Backend Files / Routes", "Database Tables"]
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        c = hrow.cells[i]
        cell_shd(c, "16227E")
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p2 = c.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(h)
        r2.font.bold  = True
        r2.font.name  = "Times New Roman"
        r2.font.size  = Pt(10)
        r2.font.color.rgb = WHITE

    for ri in range(max_rows):
        row   = t.rows[ri + 1]
        bg    = "EEF0FF" if ri % 2 == 0 else "FFFFFF"
        vals  = [frontend_rows[ri], backend_rows[ri], db_rows[ri]]
        for ci, val in enumerate(vals):
            c  = row.cells[ci]
            cell_shd(c, bg)
            p2 = c.paragraphs[0]
            r2 = p2.add_run(val)
            r2.font.name  = "Times New Roman"
            r2.font.size  = Pt(10)
            r2.font.color.rgb = DARK if val else WHITE
    doc.add_paragraph()


def screenshot_table(items):
    """
    items: list of (priority, description)
    priority: "HIGH" | "MEDIUM"
    """
    t = doc.add_table(rows=1 + len(items), cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    hrow = t.rows[0]
    for i, h in enumerate(["#", "Screenshot to Capture", "Priority"]):
        c = hrow.cells[i]
        cell_shd(c, "1A2A6C")
        p2 = c.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(h)
        r2.font.bold = True
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(10)
        r2.font.color.rgb = WHITE

    for ri, (priority, desc) in enumerate(items):
        row = t.rows[ri + 1]
        bg  = "EEF0FF" if ri % 2 == 0 else "FFFFFF"
        priority_color = "D4EDDA" if priority == "HIGH" else "FFF3CD"
        for ci, val in enumerate([str(ri + 1), desc, priority]):
            c = row.cells[ci]
            cell_shd(c, priority_color if ci == 2 else bg)
            p2 = c.paragraphs[0]
            if ci == 2:
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(val)
            r2.font.name  = "Times New Roman"
            r2.font.size  = Pt(10)
            r2.font.bold  = (ci == 2)
            r2.font.color.rgb = GREEN_TXT if priority == "HIGH" and ci == 2 else DARK
    doc.add_paragraph()


def pb():
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════════════════

title_block(
    "INSPIRABILITY — Core Features Guide",
    "Screenshot & Reference Guide for Graduation Project Documentation"
)

# Legend table
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("How to use this guide:")
r.font.name  = "Times New Roman"
r.font.size  = Pt(11)
r.font.bold  = True
r.font.color.rgb = NAVY

legend = doc.add_table(rows=1, cols=2)
legend.alignment = WD_TABLE_ALIGNMENT.CENTER
legend.style = "Table Grid"
for i, (label, meaning) in enumerate([
    ("HIGH priority screenshot", "Must have — essential evidence for the documentation"),
    ("MEDIUM priority screenshot", "Nice to have — strengthens the evidence"),
]):
    row   = legend.rows[0] if i == 0 else legend.add_row()
    c1, c2 = row.cells[0], row.cells[1]
    cell_shd(c1, "D4EDDA" if i == 0 else "FFF3CD")
    cell_shd(c2, "FFFFFF")
    p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(label)
    r1.font.bold = True; r1.font.size = Pt(10); r1.font.name = "Times New Roman"
    r1.font.color.rgb = GREEN_TXT if i == 0 else RGBColor(0x85, 0x6A, 0x04)
    r2 = c2.paragraphs[0].add_run(meaning)
    r2.font.size = Pt(10); r2.font.name = "Times New Roman"
doc.add_paragraph()

hr()

# Summary table of 7 features
section_label("CORE FEATURES COVERED IN THIS GUIDE")
features_summary = doc.add_table(rows=8, cols=3)
features_summary.style = "Table Grid"
features_summary.alignment = WD_TABLE_ALIGNMENT.CENTER
hrow = features_summary.rows[0]
for ci, h in enumerate(["#", "Feature", "Primary Purpose"]):
    c = hrow.cells[ci]; cell_shd(c, "16227E")
    p2 = c.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(h); r2.font.bold = True; r2.font.size = Pt(10)
    r2.font.name = "Times New Roman"; r2.font.color.rgb = WHITE

summary_data = [
    ("1", "Authentication System",           "User registration, login, JWT security, role-based access"),
    ("2", "Appointment Booking & Management","Core transaction: parent books, provider manages, status lifecycle"),
    ("3", "Services Directory System",       "Browse and discover schools, clinics, and sport centers"),
    ("4", "Payment & Pricing Plans",         "Flat-fee subscription plans for service providers"),
    ("5", "Admin Dashboard & Management",    "Platform analytics, provider approval, payment management"),
    ("6", "Search & Filtering",              "Filter services by location, type, specialization, price"),
    ("7", "Profile Management",              "View and edit profiles, upload media, manage account data"),
]
for ri, (num, feat, purp) in enumerate(summary_data):
    row = features_summary.rows[ri + 1]
    bg  = "EEF0FF" if ri % 2 == 0 else "FFFFFF"
    for ci, val in enumerate([num, feat, purp]):
        c = row.cells[ci]; cell_shd(c, bg)
        p2 = c.paragraphs[0]
        r2 = p2.add_run(val)
        r2.font.name = "Times New Roman"; r2.font.size = Pt(10)
        if ci == 1: r2.font.bold = True

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════
# FEATURE 1 — AUTHENTICATION
# ═══════════════════════════════════════════════════════════════════════════

pb()
feature_header(1, "Authentication System", "🔐")
hr("00A6C8")

section_label("DESCRIPTION")
body(
    "The authentication system controls access to the Inspirability platform for all user types. "
    "It handles registration (signup) for four roles — Parent, School, Medical Clinic, and Sport Center — "
    "as well as secure login using email and password. Passwords are hashed using bcrypt (salt factor 10) "
    "before storage. On successful login, the server generates a JSON Web Token (JWT) containing the "
    "user's ID, email, and role, with a 24-hour expiry. The JWT is stored in localStorage on the client "
    "and sent as a Bearer token in the Authorization header of every subsequent API request. "
    "Role-based access control (RBAC) is enforced by the authenticate() middleware on all protected routes."
)

section_label("FILE REFERENCES")
ref_table(
    frontend_rows=[
        "app/login/page.js",
        "app/signup/page.js",
        "app/components/Navbar.jsx",
    ],
    backend_rows=[
        "controllers/authController.js",
        "routes/authRoutes.js",
        "middleware: authenticate()",
    ],
    db_rows=[
        "users  (email, password, role,",
        "  approval_status,",
        "  payment_status, created_at)",
    ]
)

section_label("SUGGESTED SCREENSHOTS")
screenshot_table([
    ("HIGH",   "Login page — clean UI with email and password fields"),
    ("HIGH",   "Login page — error message for wrong password (401 response)"),
    ("HIGH",   "Login page — validation for empty fields"),
    ("HIGH",   "Signup page — full registration form for parent role"),
    ("HIGH",   "Signup page — password mismatch validation error"),
    ("HIGH",   "phpMyAdmin — users table showing hashed password + role after signup"),
    ("MEDIUM", "Signup page — role selector (showing school/clinic/sport options)"),
    ("MEDIUM", "Browser localStorage — JWT token stored after login"),
])

# ═══════════════════════════════════════════════════════════════════════════
# FEATURE 2 — APPOINTMENT BOOKING
# ═══════════════════════════════════════════════════════════════════════════

pb()
feature_header(2, "Appointment Booking & Management", "📅")
hr("00A6C8")

section_label("DESCRIPTION")
body(
    "The appointment system is the central transactional feature of the platform. A parent "
    "selects a service provider (school, clinic, or sport center), fills in a booking form "
    "with the preferred date, time, and notes, and submits the request. The booking is "
    "stored in the appointment table with a 'pending' status and linked to both the parent "
    "and the provider. The provider can then view all pending appointments in their dashboard "
    "and approve or reject each one. The parent sees the updated status in real time on their "
    "My Appointments page. Parents can cancel any pending appointment, which permanently "
    "removes the record from the database."
)

section_label("FILE REFERENCES")
ref_table(
    frontend_rows=[
        "app/appointment/page.jsx",
        "app/my-appointments/page.jsx",
        "app/provider-appointments/page.jsx",
        "app/components/ConfirmModal.jsx",
    ],
    backend_rows=[
        "controllers/appointmentController.js",
        "routes/appointmentRoutes.js",
        "  POST   /api/appointments",
        "  GET    /api/appointments/my",
        "  GET    /api/appointments/provider",
        "  PUT    /api/appointments/:id",
        "  DELETE /api/appointments/:id",
    ],
    db_rows=[
        "appointment",
        "  appointment_id",
        "  parent_id",
        "  provider_user_id",
        "  clinic_id / school_id",
        "  appointment_date",
        "  preferred_time",
        "  appointment_type",
        "  notes",
        "  status (pending/approved/rejected)",
    ]
)

section_label("SUGGESTED SCREENSHOTS")
screenshot_table([
    ("HIGH",   "Appointment booking form — parent fills date, time, and notes"),
    ("HIGH",   "My Appointments page — appointment card with PENDING status"),
    ("HIGH",   "Provider Appointments page — new pending appointment received"),
    ("HIGH",   "Provider clicks Approve — status badge changes to APPROVED"),
    ("HIGH",   "My Appointments page — appointment now shows APPROVED (green badge)"),
    ("HIGH",   "phpMyAdmin — appointment table showing the booked record"),
    ("HIGH",   "phpMyAdmin — appointment table AFTER status changed to approved"),
    ("MEDIUM", "Delete confirmation modal — before parent cancels appointment"),
    ("MEDIUM", "phpMyAdmin — appointment table AFTER deletion (row removed)"),
])

# ═══════════════════════════════════════════════════════════════════════════
# FEATURE 3 — SERVICES DIRECTORY
# ═══════════════════════════════════════════════════════════════════════════

pb()
feature_header(3, "Services Directory System", "🏫")
hr("00A6C8")

section_label("DESCRIPTION")
body(
    "The services directory is the discovery layer of the platform, allowing parents to browse "
    "all approved service providers across three verticals: Schools, Medical Clinics, and Sport "
    "Centers. Each directory page renders a searchable, filterable card grid populated by a "
    "GET request to the corresponding backend endpoint. Each provider card displays the "
    "institution's name, key attributes (location, specialization, price range), and a contact "
    "or booking prompt. Clicking a card navigates to the provider's detailed profile page. "
    "Only providers whose approval_status is 'approved' are returned by the API, ensuring "
    "quality control over directory listings."
)

section_label("FILE REFERENCES")
ref_table(
    frontend_rows=[
        "app/school/page.js",
        "app/medical/page.js",
        "app/sport/page.jsx",
        "app/services/page.js",
        "app/components/SchoolCard.jsx",
        "app/components/MedicalCard.jsx",
        "app/components/SportCard.jsx",
        "app/components/shared/FilterBar.jsx",
        "app/components/shared/CardGrid.jsx",
    ],
    backend_rows=[
        "controllers/schoolController.js",
        "controllers/medicalClinicController.js",
        "controllers/sportCenterController.js",
        "routes/schoolRoutes.js",
        "routes/medicalClinicRoutes.js",
        "routes/sportCenterRoutes.js",
    ],
    db_rows=[
        "school",
        "  school_name, curriculum_type",
        "  educational_level, city",
        "  registration_fees, annual_fees",
        "",
        "medical_clinic",
        "  clinic_name, specialization_type",
        "  session_price_range, location",
        "",
        "sport_center",
        "  sport_center_name, sports_type_offered",
        "  session_price_min/max, location",
    ]
)

section_label("SUGGESTED SCREENSHOTS")
screenshot_table([
    ("HIGH",   "Schools directory page — full grid of school cards"),
    ("HIGH",   "Medical Clinics page — full grid of clinic cards"),
    ("HIGH",   "Sport Centers page — full grid of sport center cards"),
    ("HIGH",   "Individual provider profile page — full detail view (school or clinic)"),
    ("HIGH",   "Services overview page — three service verticals displayed"),
    ("MEDIUM", "phpMyAdmin — school table showing approved provider records"),
    ("MEDIUM", "phpMyAdmin — medical_clinic table showing approved records"),
])

# ═══════════════════════════════════════════════════════════════════════════
# FEATURE 4 — PAYMENT & PRICING PLANS
# ═══════════════════════════════════════════════════════════════════════════

pb()
feature_header(4, "Payment & Pricing Plans", "💳")
hr("00A6C8")

section_label("DESCRIPTION")
body(
    "The payment system enables service providers to upgrade their subscription tier on the "
    "platform. Three pricing plans are available: Starter (free), Lite (EGP 5,000/month or "
    "EGP 4,200/year), and Pro (EGP 10,000/month or EGP 8,500/year). Providers access the "
    "flat-fee pricing page from the Navbar (visible to school, clinic, and sport roles only). "
    "After selecting a plan, they are directed to a checkout page where plan details are "
    "pre-filled. On confirmation, a payment request is stored in the users table with "
    "payment_status = 'pending'. The request appears in the admin's Requests page for "
    "manual review. The admin can approve or reject it, and the provider's account status "
    "updates accordingly."
)

section_label("FILE REFERENCES")
ref_table(
    frontend_rows=[
        "app/flat-fee/page.js",
        "app/payment/page.jsx",
        "app/payment2/page.jsx",
        "app/pricing/page.jsx",
        "app/requests/page.jsx  (admin side)",
    ],
    backend_rows=[
        "controllers/paymentController.js",
        "controllers/adminController.js",
        "  (approvePayment / rejectPayment)",
        "routes/PaymentRouter.js",
        "routes/adminRoutes.js",
        "  PUT /api/admin/approve-payment/:id",
        "  PUT /api/admin/reject-payment/:id",
        "  GET /api/admin/payment-requests",
    ],
    db_rows=[
        "users",
        "  payment_plan",
        "  payment_amount",
        "  payment_duration",
        "  payment_status",
        "  (pending / approved / rejected)",
    ]
)

section_label("SUGGESTED SCREENSHOTS")
screenshot_table([
    ("HIGH",   "Flat-fee pricing page — three plan tiers (Starter / Lite / Pro)"),
    ("HIGH",   "Flat-fee page — monthly vs yearly toggle showing price change"),
    ("HIGH",   "Checkout / payment page — plan details pre-filled from URL params"),
    ("HIGH",   "Admin Requests page — Payment Requests section with pending card"),
    ("HIGH",   "Admin approves payment — card shows Approved badge"),
    ("HIGH",   "phpMyAdmin — users table showing payment_plan, amount, status columns"),
    ("MEDIUM", "Provider navbar — Pricing link visible only for provider accounts"),
])

# ═══════════════════════════════════════════════════════════════════════════
# FEATURE 5 — ADMIN DASHBOARD
# ═══════════════════════════════════════════════════════════════════════════

pb()
feature_header(5, "Admin Dashboard & Management", "📊")
hr("00A6C8")

section_label("DESCRIPTION")
body(
    "The administrator dashboard provides the platform operator with a real-time overview of "
    "the platform's activity and full management control over provider registrations and "
    "payment requests. The dashboard renders four KPI cards (total users, total providers, "
    "new users this month, total visits) and a six-month interactive line chart comparing "
    "signups, appointments, and page visits, built with Chart.js. A growth snapshot panel "
    "shows month-over-month percentage changes. A recent-activity feed lists the five most "
    "recently registered users. The Requests page presents pending provider approval requests "
    "and payment upgrade requests, each with full provider detail and Approve/Reject actions."
)

section_label("FILE REFERENCES")
ref_table(
    frontend_rows=[
        "app/admin/page.jsx",
        "app/requests/page.jsx",
        "app/components/Navbar.jsx",
        "  (admin-specific links)",
    ],
    backend_rows=[
        "controllers/adminController.js",
        "  getStats()",
        "  getProviderRequests()",
        "  approveProvider() / rejectProvider()",
        "  getPaymentRequests()",
        "  approvePayment() / rejectPayment()",
        "  trackVisit()",
        "routes/adminRoutes.js",
    ],
    db_rows=[
        "users  (totalUsers, created_at)",
        "school",
        "medical_clinic",
        "sport_center",
        "appointment  (monthly trends)",
        "site_visits  (daily visit count)",
    ]
)

section_label("SUGGESTED SCREENSHOTS")
screenshot_table([
    ("HIGH",   "Admin Dashboard — full page view with 4 KPI cards"),
    ("HIGH",   "Admin Dashboard — six-month line chart (signups / appointments / visits)"),
    ("HIGH",   "Admin Dashboard — growth snapshot panel with percentage badges"),
    ("HIGH",   "Admin Requests page — Provider Requests section with pending provider card"),
    ("HIGH",   "Admin approves provider — card shows green Approved badge"),
    ("HIGH",   "Admin Requests page — Payment Requests section"),
    ("HIGH",   "phpMyAdmin — users table showing approval_status column values"),
    ("MEDIUM", "Admin Dashboard — recent activity feed (5 latest registrations)"),
])

# ═══════════════════════════════════════════════════════════════════════════
# FEATURE 6 — SEARCH & FILTERING
# ═══════════════════════════════════════════════════════════════════════════

pb()
feature_header(6, "Search & Filtering Functionality", "🔍")
hr("00A6C8")

section_label("DESCRIPTION")
body(
    "The search and filtering system allows parents to narrow down the service provider "
    "directories based on specific criteria relevant to their child's needs. The FilterBar "
    "component renders dynamically at the top of each directory page (Schools, Medical Clinics, "
    "Sport Centers) and provides dropdown or text filters appropriate to each service type. "
    "Filtering is performed client-side on the fetched dataset, meaning results update "
    "instantly as the parent changes filter values without any additional API calls or page "
    "reloads. This delivers a fast, responsive filtering experience."
)

section_label("FILE REFERENCES")
ref_table(
    frontend_rows=[
        "app/components/shared/FilterBar.jsx",
        "app/school/page.js",
        "app/medical/page.js",
        "app/sport/page.jsx",
    ],
    backend_rows=[
        "No dedicated filter endpoint —",
        "filtering is performed client-side",
        "on data returned by:",
        "  GET /api/school",
        "  GET /api/medical-clinic",
        "  GET /api/sport-center",
    ],
    db_rows=[
        "school",
        "  city, government",
        "  curriculum_type, educational_level",
        "",
        "medical_clinic",
        "  specialization_type, location",
        "  session_price_range",
        "",
        "sport_center",
        "  sports_type_offered",
        "  private_sessions_or_group",
        "  location",
    ]
)

section_label("SUGGESTED SCREENSHOTS")
screenshot_table([
    ("HIGH",   "Medical Clinics page — FilterBar visible with no filter applied (full results)"),
    ("HIGH",   "Medical Clinics page — filtered by specialization (e.g. Speech Therapy)"),
    ("HIGH",   "Sport Centers page — filtered by city / location"),
    ("HIGH",   "Schools page — filtered by curriculum type or educational level"),
    ("MEDIUM", "Side-by-side: unfiltered vs filtered result count for the same directory"),
])

# ═══════════════════════════════════════════════════════════════════════════
# FEATURE 7 — PROFILE MANAGEMENT
# ═══════════════════════════════════════════════════════════════════════════

pb()
feature_header(7, "Profile Management", "👤")
hr("00A6C8")

section_label("DESCRIPTION")
body(
    "The profile management system enables all registered users to view and update their "
    "account information. The profile page (/profile) renders a role-specific view using the "
    "ProfileHeader, ProfileHero, and one of four role-specific components: ParentProfile, "
    "SchoolProfile, ClinicProfile, or SportProfile. An Edit button opens the "
    "EditProfileModal — a comprehensive multi-field form that allows users to update all "
    "profile fields relevant to their role. Service providers can additionally upload and "
    "manage a media gallery (images, PDFs, certificates) through the profile page. The "
    "profile image is uploaded via multipart form data, stored on the server, and served "
    "as a static file."
)

section_label("FILE REFERENCES")
ref_table(
    frontend_rows=[
        "app/profile/page.js",
        "app/components/Profile/",
        "  ProfileHeader.jsx",
        "  ProfileHero.jsx",
        "  ParentProfile.jsx",
        "  SchoolProfile.jsx",
        "  ClinicProfile.jsx",
        "  SportProfile.jsx",
        "  EditProfileModal.jsx",
        "  ProfileInfoSections.jsx",
        "app/components/MediaGallery.jsx",
    ],
    backend_rows=[
        "routes/profileRoutes.js",
        "  GET  /api/profile",
        "  PUT  /api/profile/update",
        "  POST /api/profile/upload-image",
        "  POST /api/media/upload",
        "  GET  /api/media/:type/:id",
        "  DELETE /api/media/:mediaId",
    ],
    db_rows=[
        "parent",
        "  name, tel_no, government, city",
        "  preferred_budget, preferred_location",
        "  preferred_service_type",
        "",
        "school  /  medical_clinic",
        "  /  sport_center",
        "  (all profile detail columns)",
        "",
        "media",
        "  entity_type, entity_id",
        "  media_type, file_blob",
    ]
)

section_label("SUGGESTED SCREENSHOTS")
screenshot_table([
    ("HIGH",   "Profile page — parent profile view (name, contact, preferences)"),
    ("HIGH",   "Profile page — provider profile view (clinic or school details)"),
    ("HIGH",   "EditProfileModal — open with editable fields for provider"),
    ("HIGH",   "Profile page — after save, updated information reflected on screen"),
    ("HIGH",   "Profile page — media gallery section with uploaded images"),
    ("MEDIUM", "phpMyAdmin — parent / school / medical_clinic table after profile update"),
    ("MEDIUM", "phpMyAdmin — media table showing uploaded file records"),
])

# ═══════════════════════════════════════════════════════════════════════════
# QUICK REFERENCE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════

pb()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
para_shd(p, "16227E")
r = p.add_run("   QUICK REFERENCE — Most Important Screenshots   ")
r.font.name = "Times New Roman"
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = WHITE

doc.add_paragraph()

all_shots = [
    ("Authentication",     "Login page with valid credentials + success redirect"),
    ("Authentication",     "Login error — wrong password message"),
    ("Authentication",     "Signup form — filled out for parent role"),
    ("Authentication",     "phpMyAdmin — users table (hashed password visible)"),
    ("Appointments",       "Appointment booking form (parent side)"),
    ("Appointments",       "Provider Appointments — pending appointment card"),
    ("Appointments",       "Provider clicks Approve — badge changes"),
    ("Appointments",       "My Appointments — approved badge shown to parent"),
    ("Appointments",       "phpMyAdmin — appointment table before & after status update"),
    ("Services",           "Schools directory — full card grid"),
    ("Services",           "Medical Clinics directory — full card grid"),
    ("Services",           "Sport Centers directory — full card grid"),
    ("Services",           "Individual provider full profile page"),
    ("Payment",            "Flat-fee page — three plan tiers with toggle"),
    ("Payment",            "Checkout page — pre-filled plan details"),
    ("Payment",            "Admin Requests — Payment Requests section"),
    ("Payment",            "phpMyAdmin — users table payment columns"),
    ("Admin Dashboard",    "Full admin dashboard (KPI cards + chart)"),
    ("Admin Dashboard",    "Admin Requests — Provider Requests with Approve action"),
    ("Admin Dashboard",    "phpMyAdmin — users table (approval_status column)"),
    ("Search & Filter",    "Medical directory — filtered by specialization"),
    ("Search & Filter",    "Sport directory — filtered by location"),
    ("Profile",            "Parent profile page view"),
    ("Profile",            "Provider profile page view"),
    ("Profile",            "EditProfileModal open with editable fields"),
]

t = doc.add_table(rows=1 + len(all_shots), cols=3)
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
hrow = t.rows[0]
for ci, h in enumerate(["#", "Feature", "Screenshot Description"]):
    c = hrow.cells[ci]; cell_shd(c, "16227E")
    p2 = c.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(h)
    r2.font.bold = True; r2.font.size = Pt(10)
    r2.font.name = "Times New Roman"; r2.font.color.rgb = WHITE

feature_colors = {
    "Authentication":  "E8F4FD",
    "Appointments":    "EAF7EA",
    "Services":        "FEF9E7",
    "Payment":         "FDF2F8",
    "Admin Dashboard": "F4ECF7",
    "Search & Filter": "FDFEFE",
    "Profile":         "EBF5FB",
}

for ri, (feat, desc) in enumerate(all_shots):
    row = t.rows[ri + 1]
    bg  = feature_colors.get(feat, "FFFFFF")
    for ci, val in enumerate([str(ri + 1), feat, desc]):
        c = row.cells[ci]; cell_shd(c, bg)
        p2 = c.paragraphs[0]
        r2 = p2.add_run(val)
        r2.font.name = "Times New Roman"; r2.font.size = Pt(10)
        if ci == 1: r2.font.bold = True

doc.add_paragraph()

# closing note
note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(10)
note.paragraph_format.left_indent  = Cm(0.5)
rn = note.add_run(
    "Note:  All screenshots listed above should be captured from the running Inspirability "
    "application and inserted into the relevant chapter of the graduation report. Each "
    "screenshot must be accompanied by a figure number and a descriptive caption."
)
rn.font.name   = "Times New Roman"
rn.font.size   = Pt(11)
rn.font.italic = True
rn.font.color.rgb = GRAY_TXT

# ── Save ─────────────────────────────────────────────────────────────────────
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"Document saved: {OUTPUT_PATH}")
