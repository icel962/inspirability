"""
Inspirability User Manual Generator
Produces: Inspirability_User_Manual.docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── colour palette ────────────────────────────────────────────────────────────
PRIMARY   = RGBColor(0x16, 0x22, 0x7E)   # navy blue  #16227e
ACCENT    = RGBColor(0x00, 0xA6, 0xC8)   # teal       #00a6c8
DARK_TEXT = RGBColor(0x1A, 0x1A, 0x2E)
LIGHT_BG  = RGBColor(0xF5, 0xF7, 0xFF)

doc = Document()

# ── page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ─────────────────────────────────────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────────────────────────────────────

def set_run_font(run, bold=False, italic=False, size=11, color=None):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.font.name = "Calibri"

def add_heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_run_font(run, bold=True, size=16, color=PRIMARY)
    return p

def add_heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_run_font(run, bold=True, size=13, color=PRIMARY)
    return p

def add_heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_run_font(run, bold=True, size=11, color=ACCENT)
    return p

def add_body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    set_run_font(run, size=11)
    run.font.color.rgb = DARK_TEXT
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.8 + level * 0.6)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=11)
    run.font.color.rgb = DARK_TEXT
    return p

def add_figure_placeholder(fig_num, caption_text):
    """Insert a shaded box where a screenshot goes, plus its caption."""
    # shaded table = placeholder box
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(5.5)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "EEF0FF")
    tcPr.append(shd)
    p_inner = cell.paragraphs[0]
    p_inner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_inner.paragraph_format.space_before = Pt(28)
    p_inner.paragraph_format.space_after  = Pt(28)
    r = p_inner.add_run(f"[ Screenshot: {caption_text} ]")
    set_run_font(r, italic=True, size=10, color=RGBColor(0x80, 0x80, 0xA0))

    # caption below
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after  = Pt(14)
    cr = cap.add_run(f"Figure {fig_num} : {caption_text}")
    set_run_font(cr, italic=True, size=10, color=DARK_TEXT)
    return cap

def add_page_break():
    doc.add_page_break()

def add_horizontal_rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "16227E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ─────────────────────────────────────────────────────────────────────────────
# 0.  COVER PAGE
# ─────────────────────────────────────────────────────────────────────────────

# top spacing
doc.add_paragraph().paragraph_format.space_after = Pt(48)

# big title box
tbl = doc.add_table(rows=1, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = tbl.cell(0, 0)
tc = cell._tc
tcPr = tc.get_or_add_tcPr()
shd = OxmlElement("w:shd")
shd.set(qn("w:val"),   "clear")
shd.set(qn("w:color"), "auto")
shd.set(qn("w:fill"),  "16227E")
tcPr.append(shd)
p_title = cell.paragraphs[0]
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(20)
p_title.paragraph_format.space_after  = Pt(10)
r = p_title.add_run("User Manual")
r.font.name = "Calibri"
r.font.size = Pt(28)
r.bold = True
r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
# second line
p_sub = cell.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after  = Pt(20)
rs = p_sub.add_run("For Inspirability System")
rs.font.name = "Calibri"
rs.font.size = Pt(20)
rs.font.color.rgb = RGBColor(0xCC, 0xDD, 0xFF)

doc.add_paragraph().paragraph_format.space_after = Pt(24)

# subtitle block
def cover_line(text, sz=12, bold=False, italic=True, color=DARK_TEXT):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(sz)
    r.bold   = bold
    r.italic = italic
    r.font.color.rgb = color

cover_line("Presented to the Faculty of Business Information Systems", sz=12)
cover_line("Arab Academy for Science, Technology & Maritime Transport", sz=12)
doc.add_paragraph().paragraph_format.space_after = Pt(10)
cover_line("In Partial Fulfillment", sz=11)
cover_line("of the Requirements for the Degree of", sz=11)
cover_line("Bachelor of Business Information Systems", sz=11)
doc.add_paragraph().paragraph_format.space_after = Pt(24)

# platform name badge
p_badge = doc.add_paragraph()
p_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
rb = p_badge.add_run("INSPIRABILITY")
rb.font.name  = "Calibri"
rb.font.size  = Pt(32)
rb.bold       = True
rb.font.color.rgb = PRIMARY

doc.add_paragraph().paragraph_format.space_after = Pt(20)
cover_line("Special Needs Support Platform", sz=13, bold=True, italic=False, color=ACCENT)
doc.add_paragraph().paragraph_format.space_after = Pt(24)

# team table
team = [
    ("Supervised by", "Dr. / Project Supervisor"),
    ("Academic Year", "2025 – 2026"),
    ("Version",       "1.0"),
]
t = doc.add_table(rows=len(team), cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, value) in enumerate(team):
    t.cell(i, 0).text = label
    t.cell(i, 1).text = value
    for col in (0, 1):
        for run in t.cell(i, col).paragraphs[0].runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            if col == 0:
                run.bold = True
                run.font.color.rgb = PRIMARY

add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 1.  TABLE OF CONTENTS  (manual)
# ─────────────────────────────────────────────────────────────────────────────

add_heading1("Table of Contents")
add_horizontal_rule()

toc_entries = [
    ("1. Introduction",                                       "4"),
    ("2. System Requirements",                               "5"),
    ("   2.1  Hardware Requirements",                         "5"),
    ("   2.2  Software Requirements",                         "5"),
    ("   2.3  Mobile Access",                                 "5"),
    ("3. User Interface Navigation Guide",                   "6"),
    ("   3.1  Homepage",                                      "6"),
    ("   3.2  Sign-Up",                                       "7"),
    ("   3.3  Login",                                         "7"),
    ("   3.4  About",                                         "8"),
    ("   3.5  Services",                                      "8"),
    ("   3.6  Medical Services",                              "9"),
    ("   3.7  Sport Services",                                "9"),
    ("   3.8  Specialist Contacts (Tutors)",                 "10"),
    ("   3.9  Provider Details",                             "10"),
    ("   3.10 Pricing Plans",                                "11"),
    ("   3.11 Payment / Checkout",                           "11"),
    ("   3.12 Feedback",                                     "12"),
    ("   3.13 Contact Us",                                   "12"),
    ("   3.14 User Profile",                                 "13"),
    ("   3.15 My Appointments (Parent)",                     "13"),
    ("   3.16 Provider Appointments",                        "14"),
    ("4. Admin Interface Navigation Guide",                  "15"),
    ("   4.1  Dashboard Overview",                           "15"),
    ("   4.2  Provider Requests",                            "16"),
    ("   4.3  Payment Requests",                             "16"),
    ("   4.4  Manage Contacts",                              "17"),
    ("   4.5  Analytics & Statistics",                       "17"),
    ("5. Notes Regarding Functionalities",                   "18"),
    ("   5.1  Role-Based Access",                            "18"),
    ("   5.2  Session Management",                           "18"),
    ("   5.3  Appointment Workflow",                         "18"),
    ("   5.4  Payment Workflow",                             "18"),
]

toc_tbl = doc.add_table(rows=len(toc_entries), cols=2)
toc_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (entry, pg) in enumerate(toc_entries):
    c0 = toc_tbl.cell(i, 0)
    c1 = toc_tbl.cell(i, 1)
    c0.text = entry
    c1.text = pg
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for col, cell in ((0, c0), (1, c1)):
        para = cell.paragraphs[0]
        for run in para.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            run.font.color.rgb = DARK_TEXT
            if entry.startswith(" ") is False and col == 0:
                run.bold = True
                run.font.color.rgb = PRIMARY

add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 2.  LIST OF FIGURES  (manual)
# ─────────────────────────────────────────────────────────────────────────────

add_heading1("List of Figures")
add_horizontal_rule()

figures = [
    (1,  "Homepage – Landing Page"),
    (2,  "Sign-Up Page"),
    (3,  "Login Page"),
    (4,  "About Page"),
    (5,  "Services Overview Page"),
    (6,  "Medical Services Page"),
    (7,  "Sport Services Page"),
    (8,  "Specialist Contacts (Tutors) Page"),
    (9,  "Provider Details Page"),
    (10, "Pricing Plans Page"),
    (11, "Payment / Checkout Page"),
    (12, "Feedback Page"),
    (13, "Contact Us Page"),
    (14, "User Profile Page"),
    (15, "My Appointments – Parent View"),
    (16, "Provider Appointments Page"),
    (17, "Admin Dashboard – Overview"),
    (18, "Admin Dashboard – Provider Requests"),
    (19, "Admin Dashboard – Payment Requests"),
    (20, "Admin Dashboard – Manage Contacts"),
    (21, "Admin Dashboard – Analytics Chart"),
]

fig_tbl = doc.add_table(rows=len(figures), cols=2)
fig_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (num, caption) in enumerate(figures):
    c0 = fig_tbl.cell(i, 0)
    c1 = fig_tbl.cell(i, 1)
    c0.text = f"Figure {num} : {caption}"
    c1.text = str(num + 5)
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for cell in (c0, c1):
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            run.font.color.rgb = DARK_TEXT

add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 3.  INTRODUCTION
# ─────────────────────────────────────────────────────────────────────────────

add_heading1("1. Introduction")
add_horizontal_rule()

add_body(
    "This user manual is designed to guide users through all features and functionalities of the "
    "Inspirability System — a web-based platform developed as a graduation project for the "
    "Business Information Systems program at the Arab Academy for Science, Technology & Maritime Transport."
)

add_body(
    "Inspirability is a special needs support platform built to bridge the gap between parents of "
    "children with special needs and the professional service providers who support them. The "
    "platform brings together medical clinics, inclusive schools, adaptive sport centres, and "
    "specialist tutors in one accessible, searchable directory."
)

add_heading3("The system serves the following user groups:")
add_bullet("Parents — who wish to discover, compare, and book appointments with vetted providers.")
add_bullet("Providers — including medical clinics, schools, and sport centres that register their "
           "services and manage incoming appointment requests.")
add_bullet("Administrators — who oversee provider approvals, payment verifications, platform "
           "analytics, and contact management.")

add_body(
    "This manual provides step-by-step instructions for both end-user and administrator "
    "functionalities, including how to register, log in, browse service listings, manage profiles, "
    "book appointments, process payments, and submit feedback. The goal is to ensure that users of "
    "all technical backgrounds can comfortably navigate and make full use of the platform."
)

add_body(
    "Screenshots are included throughout the document to illustrate system behaviour and "
    "interface navigation clearly."
)

add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 4.  SYSTEM REQUIREMENTS
# ─────────────────────────────────────────────────────────────────────────────

add_heading1("2. System Requirements")
add_horizontal_rule()
add_body(
    "To ensure smooth operation and full access to all features of the Inspirability platform, "
    "users must meet the following minimum system requirements."
)

# 2.1 Hardware
add_heading2("2.1. Hardware Requirements")

hw_data = [
    ("Component",           "Minimum Requirement"),
    ("Processor",           "Dual-core 2.0 GHz or higher"),
    ("RAM",                 "4 GB or more"),
    ("Storage",             "500 MB of available space (for browser cache)"),
    ("Display",             "1280 × 720 resolution or higher"),
    ("Internet Connection", "Stable broadband with at least 5 Mbps"),
]
hw_tbl = doc.add_table(rows=len(hw_data), cols=2)
hw_tbl.style = "Table Grid"
hw_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (col1, col2) in enumerate(hw_data):
    c0 = hw_tbl.cell(i, 0)
    c1 = hw_tbl.cell(i, 1)
    c0.text = col1
    c1.text = col2
    for cell in (c0, c1):
        para = cell.paragraphs[0]
        for run in para.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"
        if i == 0:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  "16227E")
            tcPr.append(shd)

cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap.paragraph_format.space_before = Pt(4)
cap.paragraph_format.space_after  = Pt(12)
cr = cap.add_run("Table 1 : Hardware Requirements")
set_run_font(cr, italic=True, size=10)

# 2.2 Software
add_heading2("2.2. Software Requirements")

sw_data = [
    ("Category",          "Requirement"),
    ("Operating System",  "Windows 10 or later / macOS Mojave or later / Linux"),
    ("Web Browser",       "Google Chrome (v100+) / Mozilla Firefox (v95+) / Microsoft Edge"),
    ("Browser Settings",  "JavaScript enabled, cookies allowed, LocalStorage enabled"),
    ("Network",           "HTTPS-capable connection; no VPN restrictions on port 3000 / 5000"),
]
sw_tbl = doc.add_table(rows=len(sw_data), cols=2)
sw_tbl.style = "Table Grid"
sw_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (col1, col2) in enumerate(sw_data):
    c0 = sw_tbl.cell(i, 0)
    c1 = sw_tbl.cell(i, 1)
    c0.text = col1
    c1.text = col2
    for cell in (c0, c1):
        para = cell.paragraphs[0]
        for run in para.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"
        if i == 0:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  "16227E")
            tcPr.append(shd)

cap2 = doc.add_paragraph()
cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap2.paragraph_format.space_before = Pt(4)
cap2.paragraph_format.space_after  = Pt(12)
cr2 = cap2.add_run("Table 2 : Software Requirements")
set_run_font(cr2, italic=True, size=10)

# 2.3 Mobile
add_heading2("2.3. Mobile Access")
add_body(
    "The Inspirability platform is fully responsive and accessible on mobile and tablet devices "
    "via any modern mobile browser. For the best experience:"
)
add_bullet("Use the latest version of Chrome or Safari on iOS / Android.")
add_bullet("Minimum screen size: 5 inches recommended.")
add_bullet("Stable internet connection — Wi-Fi or 4G/5G recommended for media-heavy pages.")
add_bullet("Ensure JavaScript and cookies are enabled in your mobile browser settings.")

add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 5.  USER INTERFACE NAVIGATION GUIDE
# ─────────────────────────────────────────────────────────────────────────────

add_heading1("3. User Interface Navigation Guide")
add_horizontal_rule()
add_body("A step-by-step guide for parents and visitors navigating the Inspirability website.")

# ── 3.1 Homepage ─────────────────────────────────────────────────────────────
add_heading2("3.1. Homepage")
add_bullet("This is the main landing page of the Inspirability platform.")
add_bullet("The hero section features a full-screen background with the platform's tagline and "
           "two call-to-action buttons: Get Started and Explore Services.")
add_bullet("Below the hero, the Services section highlights the four main categories: "
           "Medical Clinics, Inclusive Schools, Sport Centres, and Specialist Contacts.")
add_bullet("The About / Story section describes Inspirability's mission and vision.")
add_bullet("A built-in Contact form at the bottom of the homepage allows visitors to submit "
           "general enquiries without logging in.")
add_bullet("The navigation bar provides direct links to Home, About, Services, Contact, Pricing, "
           "Login, and Sign-Up pages.")
add_figure_placeholder(1, "Homepage – Landing Page")

# ── 3.2 Sign-Up ──────────────────────────────────────────────────────────────
add_heading2("3.2. Sign-Up")
add_bullet("New users can register by clicking Sign-Up in the navigation bar.")
add_bullet("The sign-up form supports multiple roles: Parent, School, Medical Clinic, and "
           "Sport Centre.")
add_bullet("Step 1 — Select your role. Each role reveals a tailored set of registration fields.")
add_bullet("Parents fill in personal details (name, phone, government, city, child information, "
           "date of birth, preferred budget, and preferred service type).")
add_bullet("Provider accounts (School / Clinic / Sport Centre) fill in their organisation "
           "details, location, contact information, certifications, and may upload documents and a logo.")
add_bullet("All fields marked with an asterisk (*) are required. The form validates input before "
           "submission.")
add_bullet("Upon successful registration, the account is created and the user is redirected to "
           "the Login page.")
add_figure_placeholder(2, "Sign-Up Page")

# ── 3.3 Login ────────────────────────────────────────────────────────────────
add_heading2("3.3. Login")
add_bullet("Registered users can log in from the Login page accessible via the navigation bar.")
add_bullet("Enter your registered email address and password, then click Log In.")
add_bullet("The system authenticates the credentials using a secure JWT-based mechanism.")
add_bullet("Upon successful login, users are redirected based on their role:")
add_bullet("Parent accounts → redirected to the Home page.", level=1)
add_bullet("Provider accounts (School / Clinic / Sport) → redirected to their Profile page.", level=1)
add_bullet("Admin accounts → redirected to the Admin Dashboard.", level=1)
add_bullet("If credentials are incorrect, an error message is displayed. Sessions are maintained "
           "for 24 hours, after which re-login is required.")
add_figure_placeholder(3, "Login Page")

# ── 3.4 About ────────────────────────────────────────────────────────────────
add_heading2("3.4. About")
add_bullet("The About page provides in-depth information about the Inspirability platform — "
           "its mission, values, and the team behind it.")
add_bullet("Visitors can learn about the platform's commitment to supporting families of children "
           "with special needs across Egypt.")
add_bullet("Key statistics and impact metrics are displayed to build trust with new visitors.")
add_figure_placeholder(4, "About Page")

# ── 3.5 Services ─────────────────────────────────────────────────────────────
add_heading2("3.5. Services")
add_bullet("The Services overview page presents the four main service categories available on "
           "the platform.")
add_bullet("Each category card — Medical, Schools, Sports, and Specialists — includes a brief "
           "description and a Browse button that navigates to the corresponding directory.")
add_bullet("Users can click on any category to explore the full listing of registered providers.")
add_figure_placeholder(5, "Services Overview Page")

# ── 3.6 Medical Services ─────────────────────────────────────────────────────
add_heading2("3.6. Medical Services")
add_bullet("The Medical page lists all registered medical clinics and therapy centres on the "
           "platform.")
add_bullet("Each clinic card displays the clinic name, specialisation type, location, and a "
           "View Details button.")
add_bullet("Users can filter results to narrow down clinics by specialisation or location.")
add_bullet("Clicking View Details opens the provider's full profile page with their contact "
           "information, working hours, session pricing, certifications, and media gallery.")
add_figure_placeholder(6, "Medical Services Page")

# ── 3.7 Sport Services ───────────────────────────────────────────────────────
add_heading2("3.7. Sport Services")
add_bullet("The Sport page lists all registered adaptive sport centres and activity clubs.")
add_bullet("Each card shows the centre name, sport types offered, age range, and location.")
add_bullet("Filters allow users to narrow results by sport type or area.")
add_bullet("Clicking a card opens the full provider profile, which includes media uploads, "
           "session pricing, staff qualifications, and booking options.")
add_figure_placeholder(7, "Sport Services Page")

# ── 3.8 Specialist Contacts ──────────────────────────────────────────────────
add_heading2("3.8. Specialist Contacts (Tutors)")
add_bullet("The Contacts page is a searchable directory of private specialists: teachers, "
           "therapists, doctors, coaches, and tutors.")
add_bullet("Users can apply multiple filters simultaneously: specialty, budget range, distance, "
           "and rating.")
add_bullet("A live search bar allows users to search by name, email, or skill.")
add_bullet("Each contact card shows the specialist's photo, name, role, rating, and budget.")
add_bullet("Clicking a card opens the full Contact Details page, which includes a full bio, "
           "availability, and a Feedback button.")
add_figure_placeholder(8, "Specialist Contacts (Tutors) Page")

# ── 3.9 Provider Details ─────────────────────────────────────────────────────
add_heading2("3.9. Provider Details")
add_bullet("The Provider Details page (Club Details) is a dedicated profile view for each "
           "registered school, clinic, or sport centre.")
add_bullet("The page includes the provider's logo, name, full description, location, phone, "
           "email, working hours, pricing, certifications, and supported conditions.")
add_bullet("A Media Gallery section displays uploaded photos and documents associated with "
           "the provider.")
add_bullet("Four action buttons are prominently shown: Book Appointment, WhatsApp, Call, "
           "and Email — allowing parents to contact providers directly.")
add_figure_placeholder(9, "Provider Details Page")

# ── 3.10 Pricing ─────────────────────────────────────────────────────────────
add_heading2("3.10. Pricing Plans")
add_bullet("The Pricing page presents the subscription plans available to service providers who "
           "wish to list their organisation on the platform.")
add_bullet("Plans are displayed in a clean card layout with a Monthly / Yearly billing toggle at "
           "the top of the page.")
add_bullet("Three tiers are available: Starter (free), Lite (EGP 300 / month), and Pro "
           "(EGP 1,700 / month). Yearly billing offers a discounted rate.")
add_bullet("Each plan card lists its included features with tick-marks, and a Subscribe button "
           "that proceeds to the payment page.")
add_bullet("A separate Flat-Fee page exists for providers who prefer a one-off annual payment "
           "option at a higher tier (EGP 4,200 – 10,000).")
add_figure_placeholder(10, "Pricing Plans Page")

# ── 3.11 Payment ─────────────────────────────────────────────────────────────
add_heading2("3.11. Payment / Checkout")
add_bullet("The Payment page processes subscription payments for provider accounts.")
add_bullet("Three payment methods are supported: Credit / Debit Card (Visa & Mastercard), "
           "PayPal, and Apple Pay — each shown on a separate tab.")
add_bullet("The card payment form validates the cardholder name, 16-digit card number, "
           "expiry date, and CVV before submission.")
add_bullet("An Order Summary panel on the right displays the selected plan, duration, and "
           "total amount in EGP.")
add_bullet("Upon completion, the system records the payment and notifies the admin for "
           "approval. The provider's account is activated once the admin approves the payment.")
add_figure_placeholder(11, "Payment / Checkout Page")

# ── 3.12 Feedback ────────────────────────────────────────────────────────────
add_heading2("3.12. Feedback")
add_bullet("The Feedback page allows parents and users to rate their experience on the platform "
           "or with a specific provider.")
add_bullet("Users select a topic category (e.g. Appointment Experience, Platform Usability, "
           "Provider Quality) from a dropdown menu.")
add_bullet("A 1–5 star interactive rating widget captures the user's overall score.")
add_bullet("A free-text field allows the user to write a detailed comment or review.")
add_bullet("Contact fields (email and phone) are optional and can be included for follow-up.")
add_bullet("Clicking Submit sends the feedback to the platform for review.")
add_figure_placeholder(12, "Feedback Page")

# ── 3.13 Contact Us ──────────────────────────────────────────────────────────
add_heading2("3.13. Contact Us")
add_bullet("The Contact Us page provides a direct communication channel with the Inspirability "
           "support team.")
add_bullet("The page displays the platform's phone number, email address, and physical address.")
add_bullet("A Send Message form on the page includes fields for the user's name, email, and "
           "message.")
add_bullet("Submitting the form sends the message to the admin's inbox for follow-up.")
add_figure_placeholder(13, "Contact Us Page")

# ── 3.14 User Profile ────────────────────────────────────────────────────────
add_heading2("3.14. User Profile")
add_bullet("Authenticated users can access their profile by clicking Profile in the navigation bar.")
add_bullet("The profile page displays the user's registered information including name, email, "
           "phone, location, and child information (for parent accounts).")
add_bullet("An Edit Profile modal allows the user to update their personal details.")
add_bullet("Provider accounts can update their organisation information, upload a profile photo, "
           "and manage their media gallery directly from this page.")
add_bullet("A document upload section allows providers to attach certifications and documents "
           "which are displayed in their public-facing provider listing.")
add_figure_placeholder(14, "User Profile Page")

# ── 3.15 My Appointments ────────────────────────────────────────────────────
add_heading2("3.15. My Appointments (Parent)")
add_bullet("The My Appointments page is accessible to parent accounts and displays the full "
           "history of appointments they have booked.")
add_bullet("Each appointment entry shows the provider name, appointment type, scheduled date and "
           "time, notes, and current status (Pending / Approved / Rejected).")
add_bullet("A colour-coded status badge allows parents to quickly identify appointment outcomes.")
add_bullet("Appointments that are not yet approved can be cancelled by clicking the Delete icon "
           "on the appointment card.")
add_figure_placeholder(15, "My Appointments – Parent View")

# ── 3.16 Provider Appointments ──────────────────────────────────────────────
add_heading2("3.16. Provider Appointments")
add_bullet("The Provider Appointments page is accessible to registered provider accounts "
           "(School, Clinic, and Sport Centre).")
add_bullet("It displays a list of all incoming appointment requests from parent accounts, "
           "including the parent's name, requested date and time, appointment type, and notes.")
add_bullet("The provider can Accept or Reject each appointment using the action buttons. "
           "The parent's appointment status updates automatically when action is taken.")
add_bullet("This page enables providers to efficiently manage their daily scheduling and "
           "ensure capacity is not exceeded.")
add_figure_placeholder(16, "Provider Appointments Page")

add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 6.  ADMIN INTERFACE NAVIGATION GUIDE
# ─────────────────────────────────────────────────────────────────────────────

add_heading1("4. Admin Interface Navigation Guide")
add_horizontal_rule()
add_body(
    "A step-by-step guide for administrators navigating the Inspirability Admin Dashboard. "
    "The Admin Dashboard is accessible only to accounts with the 'admin' role. Logging in "
    "with an admin account automatically redirects to the dashboard at /admin."
)

# ── 4.1 Dashboard Overview ───────────────────────────────────────────────────
add_heading2("4.1. Dashboard Overview")
add_bullet("The dashboard hero section displays four live stat cards at the top of the page:")
add_bullet("Total Visitors — the total number of unique daily sessions tracked.", level=1)
add_bullet("Total Registered Users — count of all parent and provider accounts.", level=1)
add_bullet("Total Contact Listings — count of all schools, clinics, and sport centres on the platform.", level=1)
add_bullet("New This Month — count of user registrations in the current calendar month.", level=1)
add_bullet("All four values are fetched in real time from the database via the /api/admin/stats endpoint.")
add_bullet("The Insights section below the cards contains:")
add_bullet("Monthly Performance Line Chart — showing Visitors, Signups, and Appointments "
           "over the last 6 months.", level=1)
add_bullet("Growth Snapshot — showing month-on-month percentage change for each metric with "
           "colour-coded pill indicators (green = growth, red = decline).", level=1)
add_bullet("Recent Registrations — a live list of the 5 most recently created user accounts "
           "with email, role, and registration date.", level=1)
add_figure_placeholder(17, "Admin Dashboard – Overview")

# ── 4.2 Provider Requests ────────────────────────────────────────────────────
add_heading2("4.2. Provider Requests")
add_bullet("The Requests page (/requests) is accessible from the Requests link in the "
           "navigation bar (visible only to admin users).")
add_bullet("The page is split into two sections: Provider Registration Requests and "
           "Payment Requests.")
add_heading3("Provider Registration Requests")
add_bullet("When a new School, Clinic, or Sport Centre registers on the platform, their account "
           "is set to 'pending' approval status.")
add_bullet("The admin sees a card for each pending provider with their full registration "
           "details: organisation name, type, location, contact information, fees, certifications, "
           "and any uploaded documents or logo.")
add_bullet("Two action buttons are available for each request:")
add_bullet("Approve — sets the provider's approval_status to 'approved', activating their "
           "public listing.", level=1)
add_bullet("Reject — sets the status to 'rejected', preventing the provider from being "
           "listed publicly.", level=1)
add_bullet("Once actioned, the card is removed from the pending list.")
add_figure_placeholder(18, "Admin Dashboard – Provider Requests")

# ── 4.3 Payment Requests ─────────────────────────────────────────────────────
add_heading2("4.3. Payment Requests")
add_heading3("Payment Requests")
add_bullet("After a provider completes a payment for a subscription plan, their payment_status "
           "is set to 'pending'.")
add_bullet("The Payment Requests section shows all pending payments with the provider's name, "
           "role, selected plan, payment amount, and payment duration.")
add_bullet("The admin can Approve or Reject each payment:")
add_bullet("Approve — confirms the payment and fully activates the provider's subscription.", level=1)
add_bullet("Reject — declines the payment and notifies the provider to re-submit.", level=1)
add_bullet("This manual review step ensures that all payments are verified before a provider "
           "gains active status on the platform.")
add_figure_placeholder(19, "Admin Dashboard – Payment Requests")

# ── 4.4 Manage Contacts ─────────────────────────────────────────────────────
add_heading2("4.4. Manage Contacts (Specialist Directory)")
add_bullet("The Contacts Management section is located in the lower half of the Admin Dashboard.")
add_bullet("It contains a full CRUD interface for managing specialist contact listings "
           "(private teachers, doctors, coaches, and therapists).")
add_heading3("Adding a New Contact")
add_bullet("The Add New Contact form on the left accepts: Full Name, Category, Specialization, "
           "Phone, Email, Location, Availability, Status, Budget, Rating, Profile Photo, "
           "and Description.")
add_bullet("After filling in the required fields (Name and Category), click Add Contact. "
           "The entry is saved and immediately appears in the public Contacts directory.")
add_heading3("Editing a Contact")
add_bullet("Each row in the Existing Contacts table has an Edit button.")
add_bullet("Clicking Edit opens a modal with all existing values pre-filled and editable.")
add_bullet("Click Save Changes to apply updates, or Cancel to discard.")
add_heading3("Deleting a Contact")
add_bullet("Clicking the Delete button on a contact row permanently removes that specialist "
           "from the directory.")
add_heading3("Filtering Contacts")
add_bullet("The table includes live filters by: search keyword, Category, Specialty, and Status.")
add_figure_placeholder(20, "Admin Dashboard – Manage Contacts")

# ── 4.5 Analytics ───────────────────────────────────────────────────────────
add_heading2("4.5. Analytics & Statistics")
add_bullet("The analytics section provides a real-time overview of platform activity over "
           "the last 6 months.")
add_bullet("The Monthly Performance chart displays three data series:")
add_bullet("Visitors (navy line) — daily sessions tracked via the VisitTracker component.", level=1)
add_bullet("Signups (teal line) — new user registrations per month.", level=1)
add_bullet("Appointments (purple line) — appointments booked per month.", level=1)
add_bullet("The Growth Snapshot panel beneath the chart shows percentage changes month-on-month "
           "for visitors, signups, and appointments, each with a coloured pill:")
add_bullet("Green (+%) — indicates an increase compared to the previous month.", level=1)
add_bullet("Red (-%) — indicates a decrease compared to the previous month.", level=1)
add_bullet("The Recent Registrations list shows the last 5 new accounts with their email, "
           "role (parent / school / clinic / sport), and registration date.")
add_figure_placeholder(21, "Admin Dashboard – Analytics Chart")

add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 7.  NOTES REGARDING FUNCTIONALITIES
# ─────────────────────────────────────────────────────────────────────────────

add_heading1("5. Notes Regarding Functionalities")
add_horizontal_rule()

add_heading2("5.1. Role-Based Access Control")
add_body(
    "Inspirability implements strict role-based access control (RBAC). Each user account is "
    "assigned one of four roles upon registration: parent, school, clinic, sport, or admin. "
    "The navigation links displayed, the pages accessible, and the API endpoints available "
    "all depend on the authenticated role."
)
add_bullet("Parent accounts can browse providers, book appointments, manage their profile, "
           "and view their appointment history.")
add_bullet("Provider accounts (School / Clinic / Sport) can manage their profile, upload "
           "media, and view / respond to incoming appointment requests.")
add_bullet("Admin accounts have exclusive access to the Admin Dashboard, provider approval "
           "workflows, payment verification, and platform analytics.")
add_bullet("Unauthenticated visitors can browse service listings and the homepage, but cannot "
           "book appointments, submit feedback, or access profile pages.")

add_heading2("5.2. Session Management")
add_body(
    "Authentication is handled via JSON Web Tokens (JWT). Upon login, a token is issued and "
    "stored in the browser's LocalStorage. All protected API requests include this token in the "
    "Authorization header."
)
add_bullet("Sessions are valid for 24 hours. After expiry, the user is automatically redirected "
           "to the Login page on their next protected action.")
add_bullet("Clicking Logout from the navigation bar clears LocalStorage immediately and "
           "redirects the user to the Login page.")
add_bullet("Admin-only pages (/admin and /requests) detect expired or missing tokens and "
           "redirect to /login automatically.")

add_heading2("5.3. Appointment Workflow")
add_body("The appointment lifecycle follows a four-step workflow:")
add_bullet("Step 1 — Parent clicks Book Appointment on a provider's details page "
           "(/clubdetails) and is directed to the appointment form.")
add_bullet("Step 2 — The parent selects the appointment date, preferred time slot "
           "(30-minute increments, 8 AM – 10 PM), adds optional notes, and submits.")
add_bullet("Step 3 — The appointment appears in the provider's Provider Appointments page "
           "with a Pending status.")
add_bullet("Step 4 — The provider accepts or rejects the appointment. The parent sees the "
           "updated status on their My Appointments page.")
add_bullet("Appointments with a status of Approved cannot be deleted by the parent; only "
           "Pending appointments can be cancelled.")

add_heading2("5.4. Provider Activation Workflow")
add_body(
    "Before a provider's listing goes live on the platform, two sequential approval steps "
    "must be completed by the admin:"
)
add_bullet("Registration Approval — The admin reviews the provider's registration details "
           "and approves or rejects them from the Provider Requests section.")
add_bullet("Payment Approval — After the provider selects a pricing plan and submits "
           "payment, the admin verifies and approves the payment from the Payment Requests "
           "section. Only after both steps are complete is the provider's listing visible "
           "to parents on the platform.")

add_heading2("5.5. Visit Tracking")
add_body(
    "The platform includes a lightweight visitor tracking mechanism. A VisitTracker component "
    "is embedded in the root layout and runs once per browser session. It fires a POST request "
    "to /api/admin/track-visit, which increments a daily counter in the site_visits database "
    "table. This data powers the Visitors metric on the Admin Dashboard analytics chart."
)

add_heading2("5.6. Media Gallery")
add_body(
    "Provider accounts can upload images, PDFs, and documents during registration or via the "
    "Edit Profile flow. All uploaded files are stored on the server under /uploads and "
    "referenced in the media database table. The files appear in the provider's public-facing "
    "Media Gallery on their details page and in their profile's Documents section."
)

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────

out_path = r"c:\Users\Lenovo\Downloads\inspirability (2)\inspirability\my-app\Inspirability_User_Manual.docx"
doc.save(out_path)
print(f"Document saved: {out_path}")
